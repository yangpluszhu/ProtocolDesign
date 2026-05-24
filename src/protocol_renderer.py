#!/usr/bin/env python
"""Render a retrospective cohort protocol JSON file into a styled DOCX."""

from __future__ import annotations

import argparse
import io
import json
import re
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "python-docx is required. Install it with `pip install python-docx` "
        "or use an environment that already provides it."
    ) from exc


FORBIDDEN_PATTERNS = [
    "本方案模版使用说明",
    "本方案模板使用说明",
    "适用场景：",
    "注意（此部分在使用前应详细阅读",
    "【模板文本】",
    "写作要点",
    "【示例",
    "待替换占位内容",
    "正式提交前应删除",
    "【研究题目】",
    "推荐表单示例如下",
]


BODY_FONT = "宋体"
HEADING_FONT = "黑体"
LATIN_FONT = "Times New Roman"
ACCENT_COLOR = RGBColor(31, 78, 121)
LIGHT_FILL = "D9EAF7"
PALE_FILL = "F3F8FC"
BODY_FONT_SIZE = Pt(11)
TABLE_FONT_SIZE = Pt(10.5)
APPENDIX_BODY_FONT_SIZE = Pt(10)
APPENDIX_TABLE_FONT_SIZE = Pt(9)
APPENDIX_HEADING1_FONT_SIZE = Pt(14)
APPENDIX_HEADING2_FONT_SIZE = Pt(11)

import resource_crypto


def load_json(path: str) -> dict:
    if path == "-":
        return json.load(sys.stdin)
    with open(path, "r", encoding="utf-8") as handle:
        return json.load(handle)


def style_name(document: Document, preferred: str, fallback: str = "Normal") -> str:
    names = {style.name.lower(): style.name for style in document.styles}
    return names.get(preferred.lower(), names.get(fallback.lower(), fallback))


def set_run_font(run, east_asia: str = BODY_FONT, latin: str = LATIN_FONT, size: Pt | None = None) -> None:
    run.font.name = latin
    if size is not None:
        run.font.size = size
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)


def set_style_font(style, east_asia: str, latin: str = LATIN_FONT, size: Pt | None = None) -> None:
    style.font.name = latin
    if size is not None:
        style.font.size = size
    style.element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    style.element.rPr.rFonts.set(qn("w:ascii"), latin)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), latin)


def configure_page(document: Document) -> None:
    for section in document.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.6)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)
        section.different_first_page_header_footer = True


def configure_styles(document: Document) -> None:
    normal = document.styles[style_name(document, "Normal")]
    set_style_font(normal, BODY_FONT, size=BODY_FONT_SIZE)
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for preferred, font_size, before, after in [
        ("Heading 1", 16, 18, 10),
        ("heading 1", 16, 18, 10),
        ("Heading 2", 13, 12, 6),
        ("heading 2", 13, 12, 6),
    ]:
        try:
            style = document.styles[style_name(document, preferred)]
        except KeyError:
            continue
        set_style_font(style, HEADING_FONT, size=Pt(font_size))
        style.font.bold = True
        style.font.color.rgb = ACCENT_COLOR
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for style_id, font_size in [("Title", 22), ("Subtitle", 12)]:
        try:
            style = document.styles[style_name(document, style_id)]
        except KeyError:
            continue
        set_style_font(style, HEADING_FONT if style_id == "Title" else BODY_FONT, size=Pt(font_size))
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def configure_document(document: Document) -> None:
    configure_page(document)
    configure_styles(document)


def enable_field_updates_on_open(document: Document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_page_number(paragraph) -> None:
    run = paragraph.add_run("第 ")
    set_run_font(run, size=Pt(9))
    page_run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run._r.append(fld_begin)
    page_run._r.append(instr_text)
    page_run._r.append(fld_end)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=Pt(9))


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def configure_running_header_footer(document: Document, title: str) -> None:
    short_title = title if len(title) <= 34 else f"{title[:34]}..."
    for section in document.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.paragraph_format.first_line_indent = Cm(0)
        header.paragraph_format.space_after = Pt(0)
        run = header.add_run(short_title)
        set_run_font(run, size=Pt(9))
        run.font.color.rgb = RGBColor(90, 90, 90)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.paragraph_format.first_line_indent = Cm(0)
        footer.paragraph_format.space_after = Pt(0)
        add_page_number(footer)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def format_cell_text(
    cell,
    bold: bool = False,
    color: RGBColor | None = None,
    font_size: Pt = TABLE_FONT_SIZE,
) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, size=font_size)
            run.bold = bold
            if color:
                run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_key_value_table(table, label_width_cm: float = 4.2) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            set_cell_width(cell, label_width_cm if cell_index == 0 else 11.6)
            if cell_index == 0:
                set_cell_shading(cell, LIGHT_FILL if index == 0 else PALE_FILL)
                format_cell_text(cell, bold=True, color=ACCENT_COLOR)
            else:
                format_cell_text(cell)


def format_content_table(table, font_size: Pt = TABLE_FONT_SIZE) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    if not table.rows:
        return

    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            if row_index == 0:
                set_cell_shading(cell, LIGHT_FILL)
                format_cell_text(cell, bold=True, color=ACCENT_COLOR, font_size=font_size)
            elif cell_index == 0:
                set_cell_shading(cell, PALE_FILL)
                format_cell_text(cell, bold=True, color=ACCENT_COLOR, font_size=font_size)
            else:
                format_cell_text(cell, font_size=font_size)


def clear_document(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def add_paragraph(
    document: Document,
    text: str,
    style: str = "Normal",
    bold: bool = False,
    font_size: Pt | None = None,
) -> None:
    for part in split_paragraphs(text):
        paragraph = document.add_paragraph(style=style_name(document, style))
        paragraph.paragraph_format.first_line_indent = Cm(0) if style.lower().startswith("heading") else Cm(0.74)
        paragraph.paragraph_format.line_spacing = 1.5 if style.lower() == "normal" else 1.25
        run = paragraph.add_run(part)
        set_run_font(run, HEADING_FONT if style.lower().startswith("heading") else BODY_FONT, size=font_size)
        run.bold = bold


def add_content_table(
    document: Document,
    table_data,
    table_style: str | None = None,
    font_size: Pt = TABLE_FONT_SIZE,
) -> None:
    if isinstance(table_data, dict):
        title = table_data.get("title")
        headers = table_data.get("headers") or []
        rows = table_data.get("rows") or []
    else:
        title = None
        headers = []
        rows = table_data or []

    if title:
        add_paragraph(document, str(title), bold=True, font_size=font_size)

    normalized_rows = [[str(cell) for cell in row] for row in rows if row]
    normalized_headers = [str(cell) for cell in headers]
    column_count = max(
        [len(normalized_headers), *(len(row) for row in normalized_rows)] or [0]
    )
    if column_count == 0:
        return

    table = document.add_table(rows=0, cols=column_count)
    for candidate in [table_style, "Table Grid"]:
        if not candidate:
            continue
        try:
            table.style = candidate
            break
        except Exception:
            continue

    if normalized_headers:
        cells = table.add_row().cells
        for index in range(column_count):
            cells[index].text = normalized_headers[index] if index < len(normalized_headers) else ""

    for row in normalized_rows:
        cells = table.add_row().cells
        for index in range(column_count):
            cells[index].text = row[index] if index < len(row) else ""

    format_content_table(table, font_size=font_size)
    document.add_paragraph()


def split_paragraphs(text: str) -> list[str]:
    if not text:
        return []
    parts = [part.strip() for part in re.split(r"\n\s*\n", str(text).strip())]
    return [part for part in parts if part]


def add_heading(document: Document, text: str, level: int, font_size: Pt | None = None) -> None:
    style = "heading 1" if int(level) == 1 else "heading 2"
    add_paragraph(document, text, style=style, font_size=font_size)


def add_table_of_contents(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Cm(0)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("目录")
    set_run_font(run, HEADING_FONT, size=Pt(16))
    run.bold = True
    run.font.color.rgb = ACCENT_COLOR

    toc = document.add_paragraph()
    toc.paragraph_format.first_line_indent = Cm(0)
    toc.paragraph_format.line_spacing = 1.5
    add_field(toc, r'TOC \o "1-2" \h \z \u')
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def is_appendix_heading(heading: str) -> bool:
    normalized = heading.strip()
    return normalized.startswith("附录") or normalized.startswith("附表")


def add_cover(document: Document, cover: dict) -> None:
    add_heading(document, "方案封面", 1)
    title = cover.get("研究题目") or cover.get("题目") or "回顾性队列研究方案"
    document.add_paragraph()
    paragraph = document.add_paragraph(style=style_name(document, "Title"))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    set_run_font(run, HEADING_FONT, size=Pt(22))
    run.bold = True
    run.font.color.rgb = ACCENT_COLOR
    subtitle = document.add_paragraph(style=style_name(document, "Subtitle"))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("回顾性队列研究方案")
    set_run_font(subtitle_run, BODY_FONT, size=Pt(12))

    fields = [
        "研究发起单位",
        "牵头研究单位",
        "主要研究者",
        "参与单位/中心",
        "统计分析负责人",
        "数据管理负责人",
        "方案编号",
        "版本日期",
        "版本号",
    ]
    defaults = {
        "研究发起单位": "研究发起单位待研究团队确认",
        "牵头研究单位": "牵头研究单位待研究团队确认",
        "主要研究者": "主要研究者待研究团队确认",
        "参与单位/中心": "参与单位或中心待研究团队确认",
        "统计分析负责人": "统计分析负责人待研究团队确认",
        "数据管理负责人": "数据管理负责人待研究团队确认",
        "方案编号": "待编号",
        "版本日期": date.today().isoformat(),
        "版本号": "V1.0",
    }
    document.add_paragraph()
    table = document.add_table(rows=0, cols=2)
    for field in fields:
        value = cover.get(field) or defaults[field]
        cells = table.add_row().cells
        cells[0].text = field
        cells[1].text = str(value)
    format_key_value_table(table, label_width_cm=4.0)

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_summary_table(document: Document, rows: list[list[str]], table_style: str | None = None) -> None:
    add_heading(document, "方案摘要", 1)
    table = document.add_table(rows=0, cols=2)
    for candidate in [table_style, "Table Grid"]:
        if not candidate:
            continue
        try:
            table.style = candidate
            break
        except Exception:
            continue

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)
    format_key_value_table(table)
    document.add_paragraph()


def normalize_summary_rows(data: dict) -> list[list[str]]:
    supplied = data.get("summary") or []
    if isinstance(supplied, dict):
        supplied = [[key, value] for key, value in supplied.items()]
    return [[str(row[0]), str(row[1])] for row in supplied if len(row) >= 2]


def add_sections(document: Document, sections: list[dict], table_style: str | None = None) -> None:
    in_appendix = False
    for section in sections:
        heading = section.get("heading")
        if heading and is_appendix_heading(str(heading).replace("（若适用）", "")):
            in_appendix = True
        if heading:
            level = int(section.get("level", 1))
            appendix_heading_size = (
                APPENDIX_HEADING1_FONT_SIZE if level == 1 else APPENDIX_HEADING2_FONT_SIZE
            ) if in_appendix else None
            add_heading(
                document,
                str(heading).replace("（若适用）", "").strip(),
                level,
                appendix_heading_size,
            )
        for paragraph in section.get("paragraphs", []):
            add_paragraph(
                document,
                str(paragraph),
                font_size=APPENDIX_BODY_FONT_SIZE if in_appendix else None,
            )
        for table in section.get("tables", []):
            add_content_table(
                document,
                table,
                table_style,
                font_size=APPENDIX_TABLE_FONT_SIZE if in_appendix else TABLE_FONT_SIZE,
            )


def all_text(document: Document) -> str:
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def validate(document: Document) -> None:
    text = all_text(document)
    offenders = [pattern for pattern in FORBIDDEN_PATTERNS if pattern in text]
    bracket_placeholders = re.findall(r"【[^】]{1,40}】", text)
    if bracket_placeholders:
        offenders.extend(sorted(set(bracket_placeholders))[:10])
    square_placeholders = re.findall(r"\[[^\]\r\n]{0,40}[\u4e00-\u9fff][^\]\r\n]{0,40}\]", text)
    if square_placeholders:
        offenders.extend(sorted(set(square_placeholders))[:10])
    if offenders:
        formatted = "\n".join(f"- {item}" for item in offenders)
        raise SystemExit(f"Output still contains template artifacts:\n{formatted}")


def render_protocol(data: dict, output: Path, template: io.BytesIO | Path | None = None) -> Path:
    """Render protocol JSON data to DOCX and return the output path."""
    template_source = template or resource_crypto.decrypt_resource_bytesio("assets/corhortCRU-small.docx")
    document = Document(template_source)
    configure_document(document)
    enable_field_updates_on_open(document)
    title = data.get("cover", {}).get("研究题目") or data.get("cover", {}).get("题目") or "回顾性队列研究方案"
    configure_running_header_footer(document, str(title))
    table_style = document.tables[0].style.name if document.tables else None
    clear_document(document)
    add_cover(document, data.get("cover", {}))
    add_table_of_contents(document)
    add_summary_table(document, normalize_summary_rows(data), table_style)
    add_sections(document, data.get("sections", []), table_style)
    validate(document)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    return output


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("json", help="Protocol content JSON path, or '-' for stdin")
    parser.add_argument(
        "--template",
        type=Path,
        default=None,
        help="Template DOCX path; defaults to the bundled encrypted template",
    )
    parser.add_argument("--output", type=Path, default=Path("cohortprotocol.docx"), help="Output DOCX path")
    args = parser.parse_args()

    data = load_json(args.json)
    output = render_protocol(data, args.output, args.template)
    print(str(output))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
