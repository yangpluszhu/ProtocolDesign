from __future__ import annotations

from datetime import datetime
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "ProtocolDesign_帮助文档.docx"
SHOT = ROOT / "help_assets" / "ProtocolDesign_主界面.png"
ANNOTATED = ROOT / "help_assets" / "ProtocolDesign_主界面_标注.png"

PRODUCER = "上海中医药大学附属龙华医院临床研究中心"
APP_NAME_EN = "ProtocolDesign"
APP_NAME_CN = "临床研究方案辅助构建系统"
APP_NAME = f"{APP_NAME_CN}（{APP_NAME_EN}）"
EMAIL = "yangpluszhu@sina.com"
GITHUB = "https://github.com/yangpluszhu/ProtocolDesign"
VERSION = "V1.0"
THINKING = "思考模式：DeepSeek、GLM、MiMo 默认开启并请求最大强度；Kimi 按当前版本要求禁用思考模式。"
TARGET_PROTOCOL_CHINESE_CHARS = 10000
SCOPE = "适用范围：目前版本仅开发了“回顾性队列研究方案”模块，仅适用于回顾性队列研究的方案设计与撰写。其他模块为预留待开发功能。"
RISK = (
    "请注意：本方案由AI工具辅助生成，仅供研究团队参考，不应被视为最终结论或决策依据。"
    "AI生成内容可能存在事实不准确、逻辑推断不充分、引用来源不完整或过时、对特定研究场景理解偏差等风险，"
    "也可能遗漏关键变量、伦理合规要求、数据安全与知识产权风险。研究团队在采纳前应结合专业判断，"
    "对方案中的研究假设、技术路线、数据来源、试验设计、合规要求及可行性进行逐项核验，并由相关领域专家进行审查确认。"
    "任何基于本方案开展的研究、申报或实施工作，均应以人工复核和正式论证结果为准，"
    "软件制作方不对由该软件生成的结果产生的任何不良结果负责。"
)


def font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    for path in candidates:
        if not Path(path).exists():
            continue
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


def make_annotated_image() -> None:
    if not SHOT.exists():
        raise FileNotFoundError(f"主界面截图不存在：{SHOT}")
    with Image.open(SHOT) as source:
        img = source.convert("RGB")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 990, 126), fill=(255, 255, 255))
    title_size = 30
    title_font = font(title_size)
    while title_size > 18 and draw.textbbox((0, 0), APP_NAME, font=title_font)[2] > 940:
        title_size -= 2
        title_font = font(title_size)
    draw.text((20, 24), APP_NAME, font=title_font, fill=(31, 78, 121))
    draw.text((20, 86), APP_NAME_EN, font=font(18), fill=(72, 101, 117))
    callouts = [
        (1, 0.3678, 0.1558, "选择功能模块"),
        (2, 0.3079, 0.2299, "选择摘要文档"),
        (3, 0.3678, 0.3499, "选择模型类型"),
        (4, 0.3678, 0.5146, "填写/保存 API Key"),
        (5, 0.3678, 0.5833, "选择或输入模型"),
        (6, 0.3079, 0.7027, "选择输出目录"),
        (7, 0.3678, 0.7570, "填写输出文件名"),
        (8, 0.0654, 0.8167, "点击生成"),
        (9, 0.4952, 0.9373, "查看运行状态"),
    ]
    f_num = font(24)
    f_text = font(18)
    for idx, x_ratio, y_ratio, label in callouts:
        x = int(img.width * x_ratio)
        y = int(img.height * y_ratio)
        r = 20
        draw.ellipse((x - r, y - r, x + r, y + r), fill=(204, 46, 30), outline=(255, 255, 255), width=3)
        text = str(idx)
        bbox = draw.textbbox((0, 0), text, font=f_num)
        draw.text((x - (bbox[2] - bbox[0]) / 2, y - (bbox[3] - bbox[1]) / 2 - 2), text, font=f_num, fill="white")
        lx = min(x + 28, img.width - 210)
        ly = max(8, y - 18)
        tw = draw.textlength(label, font=f_text)
        draw.rounded_rectangle((lx - 6, ly - 4, lx + tw + 8, ly + 25), radius=6, fill=(255, 248, 232), outline=(204, 46, 30), width=2)
        draw.text((lx, ly), label, font=f_text, fill=(122, 41, 20))
    ANNOTATED.parent.mkdir(parents=True, exist_ok=True)
    img.save(ANNOTATED)


def set_run_font(run, east="微软雅黑", latin="Arial", size=None) -> None:
    run.font.name = latin
    if size:
        run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)


def set_style(style, east, latin="Arial", size=10.5, color=None, bold=False) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor(*color)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), east)
    style.element.rPr.rFonts.set(qn("w:ascii"), latin)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), latin)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, end])


def setup_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.2)
    normal = doc.styles["Normal"]
    set_style(normal, "微软雅黑", size=10.5)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size, color in [
        ("Title", 24, (15, 76, 107)),
        ("Heading 1", 16, (15, 76, 107)),
        ("Heading 2", 13, (31, 78, 121)),
        ("Heading 3", 11.5, (83, 83, 83)),
    ]:
        set_style(doc.styles[style_name], "微软雅黑", size=size, color=color, bold=True)
    for sec in doc.sections:
        header = sec.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run(f"{APP_NAME} 帮助文档 | {VERSION}")
        set_run_font(run, size=9)
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(footer.add_run("第 "), size=9)
        add_field(footer, "PAGE")
        set_run_font(footer.add_run(" 页"), size=9)


def p(doc, text="", style=None, align=None, bold=False, color=None, size=None):
    para = doc.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    set_run_font(run, size=size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para


def heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        set_run_font(run, size=16 if level == 1 else 13)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    set_run_font(run)


def numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    set_run_font(run)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for r, row in enumerate(t.rows):
        for c, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r == 0:
                shade_cell(cell, "D9EAF7")
            elif c == 0:
                shade_cell(cell, "F3F8FC")
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    set_run_font(run, size=9.5)
                    if r == 0 or c == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 78, 121)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return t


def cover(doc: Document) -> None:
    p(doc, APP_NAME, "Title", WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, APP_NAME_EN, None, WD_ALIGN_PARAGRAPH.CENTER, size=12, color=(72, 101, 117))
    p(doc, "用户帮助文档", None, WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18)
    doc.add_paragraph()
    table(
        doc,
        ["项目", "内容"],
        [
            ["软件制作方", PRODUCER],
            ["版本号", VERSION],
            ["作者联系方式", f"Email: {EMAIL}\nGitHub: {GITHUB}"],
            ["文档生成时间", datetime.now().strftime("%Y-%m-%d %H:%M")],
        ],
        [4.2, 11.0],
    )
    doc.add_paragraph()
    p(doc, "本文档面向首次使用者编写，按照“准备材料、填写界面、生成文档、检查结果”的顺序说明。即使没有编程经验，也可以依照步骤完成操作。")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_toc(doc: Document) -> None:
    heading(doc, "目录", 1)
    para = doc.add_paragraph()
    add_field(para, r'TOC \o "1-3" \h \z \u')
    p(doc, "提示：首次在 Word 中打开后，如目录页码未显示，请右键目录区域并选择“更新域”或“更新整个目录”。", color=(120, 70, 20))
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_screenshot(doc: Document) -> None:
    heading(doc, "一、认识主界面", 1)
    p(doc, "启动软件后，可以看到如下主界面。红色编号对应后续操作步骤。")
    p(doc, SCOPE, color=(15, 76, 107))
    doc.add_picture(str(ANNOTATED), width=Cm(16.4))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p(doc, f"图 1  {APP_NAME} 主界面及关键操作区域", align=WD_ALIGN_PARAGRAPH.CENTER, size=9, color=(90, 90, 90))
    table(
        doc,
        ["编号", "界面区域", "作用"],
        [
            ["1", "功能模块", "当前仅“回顾性队列研究方案”可用；其余 5 个模块为待开发预留。"],
            ["2", "方案摘要文档", "选择研究团队已经准备好的摘要文件，支持 docx、md、txt。"],
            ["3", "大模型类型/思考模式", f"选择 deepseek、kimi、GLM、mimo 或自定义；软件会按不同提供商自动适配思考模式。{THINKING}"],
            ["4", "api_key", "填写模型平台提供的密钥；默认隐藏。可勾选“本地保存 api_key”，软件会使用 Windows 当前用户 DPAPI 加密保存。"],
            ["5", "模型名称", "可从下拉框选择常见模型，也可手动输入平台实际支持的任意模型名。"],
            ["6", "输出目录", "选择生成 DOCX 和运行简报保存的位置。"],
            ["7", "输出文件名", "填写生成方案的文件名；不写 .docx 时软件会自动补齐。"],
            ["8", "确定生成方案", "开始读取摘要、调用模型、生成完整研究方案。"],
            ["9", "运行状态与完成简报", "显示当前进度、错误信息和完成后的输出路径。"],
        ],
        [1.5, 4.2, 10.2],
    )


def add_steps(doc: Document) -> None:
    heading(doc, "二、开始前需要准备什么", 1)
    bullet(doc, "一份方案摘要文档：建议使用 Word 的 .docx，也可以使用 .md 或 .txt。")
    bullet(doc, "一个大模型平台账号：deepseek、kimi、GLM、mimo 或用户自定义提供商，只要提供 OpenAI 兼容接口即可。")
    bullet(doc, "模型接口信息：base_url、api_key、模型名称。")
    bullet(doc, "一个输出目录：建议新建一个专门文件夹保存方案和运行简报。")
    bullet(doc, "稳定网络环境：生成完整研究方案通常需要等待数分钟，取决于模型速度和摘要复杂度。")

    heading(doc, "三、9 步完成第一次生成", 1)
    steps = [
        ("第 1 步：打开软件", "双击 app 文件夹中的 ProtocolDesign.exe。若 Windows 出现安全提醒，请确认文件来源为本项目目录后选择继续运行。"),
        ("第 2 步：选择功能模块", "在“功能模块”中选择“回顾性队列研究方案”。当前版本仅该模块可运行，其他 5 个模块为后续升级预留。"),
        ("第 3 步：选择方案摘要文档", "点击“方案摘要文档”右侧的“浏览...”按钮，选择研究团队准备好的摘要文件。软件支持 .docx、.md、.txt。建议摘要中写清楚研究题目、人群、暴露、对照、结局、统计分析设想和伦理要求。"),
        ("第 4 步：选择大模型类型", f"在“大模型类型”下拉框中选择 deepseek、kimi、GLM、mimo 或自定义。选择内置提供商会自动填入常见 base_url、候选模型名和对应调用参数；选择“自定义”时，base_url、api_key 和模型名称均由用户自行填写。{THINKING}"),
        ("第 5 步：填写模型接口信息", "依次填写或确认 base_url、api_key 和模型名称。base_url 是 OpenAI 兼容接口地址；api_key 是平台访问密钥，可按需勾选“显示 api_key”或“本地保存 api_key”；模型名称可从下拉框选择，也可手动输入平台实际支持的模型名。"),
        ("第 6 步：选择输出目录", "点击“输出目录”右侧的“浏览...”按钮，选择保存生成结果的位置。建议路径不要过深，也尽量避免使用特殊符号。"),
        ("第 7 步：填写输出文件名", "例如：益气养阴方治疗肺癌回顾性队列研究方案.docx。没有填写 .docx 后缀时，软件会自动补齐。"),
        ("第 8 步：点击确定生成方案", "点击“确定生成方案”。运行期间不要关闭软件。界面下方会显示读取摘要、调用模型、解析 JSON、渲染 DOCX 等进度。"),
        ("第 9 步：查看结果并人工复核", "完成后，输出目录中会出现研究方案 DOCX 和同名的运行简报 txt；如生成失败，程序目录会额外写出 ProtocolDesign_error.log 供排查。请打开 Word 文档检查目录、章节、附表和所有关键方法学内容，并按研究团队要求进行人工复核。"),
    ]
    for title, body in steps:
        heading(doc, title, 2)
        p(doc, body)


def add_interface_examples(doc: Document) -> None:
    heading(doc, "四、模型接口如何填写", 1)
    p(doc, "不同服务商的命名可能随平台更新而变化，请以服务商控制台或官方文档为准。下表给出常见填写方式，仅作为入门示例。")
    table(
        doc,
        ["模型类型", "base_url 示例", "模型名示例", "说明"],
        [
            ["deepseek", "https://api.deepseek.com/v1", "deepseek-v4-pro", "默认发送 thinking.type=enabled 与 reasoning_effort=max。"],
            ["kimi", "https://api.moonshot.cn/v1", "kimi-k2.6 / kimi-k2.5", "自动使用 temperature=0.6、禁用思考模式并设置较大的输出 token 预算；对 Kimi 采用封面摘要、章节分批生成后本地组装的流程；流式响应强制按 UTF-8 解码；封面摘要阶段如未返回最终正文，会继续尝试 JSON 模式和非流式禁用思考模式；章节分批阶段会继续尝试非流式禁用思考模式。"],
            ["GLM", "https://open.bigmodel.cn/api/paas/v4", "glm-4.5 / glm-4.7 / glm-5-turbo", "默认发送 thinking.type=enabled，并尝试最大 effort。"],
            ["mimo", "https://api.xiaomimimo.com/v1", "mimo-v2.5-pro / mimo-v2.5 / mimo-v2-flash", "默认尝试 thinking/reasoning 最大强度参数；若返回 JSON 含未转义控制字符，会自动修复解析，必要时按严格 JSON 重试一次。"],
            ["自定义", "用户自行填写", "用户自行填写", "适合其他 OpenAI 兼容接口；base_url、api_key 和模型名称均需用户确认。"],
        ],
        [2.2, 5.1, 3.3, 5.0],
    )
    heading(doc, "五、api_key 本地保存与安全说明", 1)
    p(doc, "软件默认不强制保存 api_key。是否保存由用户通过复选框决定。")
    for item in [
        "勾选“本地保存 api_key”后，软件会按当前模型提供商分别保存密钥，例如 deepseek、kimi、GLM、mimo 分开保存。",
        "保存位置为当前 Windows 用户的 AppData 目录：%APPDATA%\\ProtocolDesign\\settings.json。",
        "保存内容不是明文 api_key，而是通过 Windows DPAPI 以当前登录用户身份加密后的密文。",
        "同一台电脑上，其他 Windows 用户通常无法解密该密文；复制到其他电脑后也通常无法解密。",
        "取消勾选“本地保存 api_key”后，软件会删除当前提供商对应的本地保存项。",
        "运行简报、错误日志和帮助文档不会记录 api_key。",
        "如电脑属于多人共用设备，建议不要勾选保存，并在使用后关闭软件。",
    ]:
        bullet(doc, item)

    heading(doc, "六、摘要文档怎么写更容易成功", 1)
    p(doc, "摘要不是越长越好，而是要包含关键研究设计信息。建议至少包括以下内容：")
    for item in [
        "研究题目：尽量写成完整医学研究题目。",
        "研究背景与目的：说明为什么要做、主要想回答什么问题。",
        "数据来源：医院电子病历、登记数据库、随访数据库、医保数据等。",
        "研究对象：目标人群、纳入标准、排除标准、研究时间范围。",
        "暴露与对照：暴露组怎么定义，对照组怎么定义，索引日期如何确定。",
        "结局指标：主要结局、次要结局、安全性结局及识别方式。",
        "协变量：年龄、性别、疾病严重程度、合并症、合并用药、既往治疗等。",
        "统计分析：倾向评分、Cox 模型、竞争风险、亚组分析、敏感性分析等初步设想。",
        "伦理与数据安全：伦理审查、知情同意豁免、去标识化、数据访问权限等。",
    ]:
        bullet(doc, item)


def add_outputs_and_troubleshooting(doc: Document) -> None:
    heading(doc, "七、生成后会得到哪些文件", 1)
    p(doc, f"为了保证方案内容完整详尽，软件会在提示词中要求模型参照方案模板示例的写作风格和写作要点尽量详尽撰写，整体中文字符数建议不少于 {TARGET_PROTOCOL_CHINESE_CHARS} 个。该长度要求仅作为生成目标，不作为输出限制；即使最终文档未达到 {TARGET_PROTOCOL_CHINESE_CHARS} 个中文字符，软件也会放行输出。")
    table(
        doc,
        ["文件", "说明", "建议操作"],
        [
            ["完整研究方案 .docx", "软件根据摘要和模板生成的正式方案文档。", "使用 Word 打开，先更新目录，再进行专家复核。"],
            ["运行简报 .txt", "记录输入文件、输出文件、模型名称、耗时、章节数量和风险提示。", "与方案一起保存，便于追溯生成过程。"],
            ["错误日志 ProtocolDesign_error.log", "仅当生成失败时出现，记录详细错误信息。", "根据错误提示检查接口、网络、模型输出或摘要文件。"],
        ],
        [4.2, 6.2, 5.0],
    )

    heading(doc, "八、常见问题与处理办法", 1)
    table(
        doc,
        ["现象", "常见原因", "处理办法"],
        [
            ["提示 base_url 或 api_key 错误", "接口地址、密钥、模型名填写错误，或账号额度不足。", "重新复制平台控制台信息；确认账号可用；检查 base_url 是否为 OpenAI 兼容地址。"],
            ["模型返回 400 错误", "模型名不存在、response_format 或思考参数不被支持、请求格式被平台拒绝。Kimi 请求已自动使用 temperature=0.6 并禁用思考模式。", "改用平台实际模型名；软件会按不同提供商自动尝试兼容参数和普通 JSON 模式。"],
            ["Kimi 返回 429 或 engine_overloaded", "模型服务临时过载或排队，接口提示稍后重试。", "软件会自动等待并重试；若多次重试后仍失败，请稍后再试或临时切换模型。"],
            ["Kimi 输出 JSON 不完整", "模型返回了 JSON，但片段中出现字符串未闭合、语法不完整，或某个章节批次缺少指定标题。", "软件会对 Kimi 采用分段生成：先生成封面摘要，再按章节批次生成 sections；某个片段失败时只重试当前片段，最终由软件本地组装完整方案。"],
            ["MiMo JSON 解析失败", "模型返回的 JSON 字符串中可能包含未转义换行、制表符等控制字符。", "软件会自动修复后解析；如仍失败，会按严格 JSON 要求重试一次。"],
            ["输出中文显示乱码", "流式接口响应头未正确声明编码，客户端按错误编码解析了中文。", "软件会对流式响应强制使用 UTF-8，并对明显乱码文本做保守修复，正常中文不会被改动。"],
            ["生成内容偏短", "模型未充分展开模板写作要点，或当前模型输出长度能力有限。", "软件不会因字符数阻断输出；如需更详尽版本，可重新生成或改用上下文更长、输出能力更强的模型。"],
            ["提示模型输出不是 JSON", "模型没有严格按照要求输出结构化 JSON。", "换用更强模型，或让摘要更清晰后重新生成。"],
            ["生成时间很长", "完整方案内容多，模型推理和输出耗时较长。", "保持网络稳定，不要关闭软件；可查看下方运行状态。"],
            ["Word 目录没有页码", "Word 尚未更新目录域。", "右键目录区域，选择“更新域”或“更新整个目录”。"],
            ["生成内容需要修改", "AI 输出仍需人工复核和本地化调整。", "由研究者、统计师、伦理和数据管理人员逐项审查后修订。"],
        ],
        [4.0, 5.4, 6.0],
    )

    heading(doc, "九、研究团队人工复核清单", 1)
    for item in [
        "研究假设是否清晰，是否与摘要和实际研究目的一致。",
        "暴露、对照和 time zero 是否定义明确，是否避免 immortal-time bias。",
        "纳入排除标准是否能在数据源中被真实识别。",
        "主要结局、次要结局和安全性结局是否有可执行的编码或算法。",
        "基线协变量是否均在 time zero 前测量。",
        "混杂控制、缺失数据处理、亚组分析和敏感性分析是否合理。",
        "伦理审查、知情同意豁免、隐私保护和数据安全表述是否符合本单位要求。",
        "附表中的操作性定义是否能直接支持数据提取和统计分析。",
        "所有推断性内容是否已由相关领域专家确认。",
    ]:
        bullet(doc, item)

    heading(doc, "十、风险提示", 1)
    p(doc, RISK, color=(122, 41, 20))


def main() -> None:
    try:
        make_annotated_image()
        doc = Document()
        setup_doc(doc)
        cover(doc)
        add_toc(doc)
        add_screenshot(doc)
        add_steps(doc)
        add_interface_examples(doc)
        add_outputs_and_troubleshooting(doc)
        doc.save(OUT)
        print(OUT)
    except Exception as exc:
        raise SystemExit(f"帮助文档生成失败：{exc}") from exc


if __name__ == "__main__":
    main()
