from __future__ import annotations

import json
import os
import re
import sys
import threading
import time
import traceback
import base64
import ctypes
from ctypes import wintypes
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from tkinter import BOTH, END, LEFT, RIGHT, X, Y, BooleanVar, StringVar, TclError, Tk, Canvas, Text, filedialog, messagebox
from tkinter import ttk

import requests
from docx import Document

import protocol_renderer
import resource_crypto


APP_NAME_EN = "ProtocolDesign"
APP_NAME_CN = "研案智构——临床研究方案辅助构建系统"
APP_NAME = f"{APP_NAME_CN}（{APP_NAME_EN}）"
APP_VERSION = "V1.0"
PRODUCER = "上海中医药大学附属龙华医院临床研究中心"
AUTHOR_EMAIL = "yangpluszhu@sina.com"
AUTHOR_GITHUB = "https://github.com/yangpluszhu/ProtocolDesign"
RISK_NOTICE = (
    "请注意：本方案由AI工具辅助生成，仅供研究团队参考，不应被视为最终结论或决策依据。"
    "AI生成内容可能存在事实不准确、逻辑推断不充分、引用来源不完整或过时、对特定研究场景理解偏差等风险，"
    "也可能遗漏关键变量、伦理合规要求、数据安全与知识产权风险。研究团队在采纳前应结合专业判断，"
    "对方案中的研究假设、技术路线、数据来源、试验设计、合规要求及可行性进行逐项核验，并由相关领域专家进行审查确认。"
    "任何基于本方案开展的研究、申报或实施工作，均应以人工复核和正式论证结果为准，"
    "软件制作方不对由该软件生成的结果产生的任何不良结果负责。"
)

PROVIDER_DEFAULTS = {
    "deepseek": {"base_url": "https://api.deepseek.com/v1", "model": "deepseek-v4-pro"},
    "kimi": {"base_url": "https://api.moonshot.cn/v1", "model": "kimi-k2.6"},
    "GLM": {"base_url": "https://open.bigmodel.cn/api/paas/v4", "model": "glm-4.5"},
    "mimo": {"base_url": "https://api.xiaomimimo.com/v1", "model": "mimo-v2.5-pro"},
    "自定义": {"base_url": "", "model": ""},
}

PROVIDER_MODELS = {
    "deepseek": ["deepseek-v4-pro", "deepseek-v4-flash", "deepseek-reasoner"],
    "kimi": ["kimi-k2.6", "kimi-k2.5", "kimi-k2-thinking", "kimi-k2-thinking-turbo"],
    "GLM": ["glm-4.5", "glm-4.5-air", "glm-4.7", "glm-4.7-flash", "glm-5", "glm-5.1", "glm-5-turbo"],
    "mimo": ["mimo-v2.5-pro", "mimo-v2.5", "mimo-v2-pro", "mimo-v2-omni", "mimo-v2-flash"],
    "自定义": [],
}

THINKING_MODE_NOTICE = (
    "思考模式：DeepSeek、GLM、MiMo 默认开启并请求最大强度；Kimi 按当前版本要求禁用思考模式。"
)
SCOPE_NOTICE = "适用范围：目前版本仅开发了“回顾性队列研究方案”模块，仅适用于回顾性队列研究的方案设计与撰写。其他模块为预留待开发功能。"
TARGET_PROTOCOL_CHINESE_CHARS = 10000
TRANSIENT_STATUS_CODES = {429, 500, 502, 503, 504}
TRANSIENT_RETRY_DELAYS = [20, 40, 80, 120, 240]
MODEL_REQUEST_TIMEOUT_SECONDS = 5400
KIMI_VALIDATION_RETRY_LIMIT = 2
KIMI_SEGMENT_RETRY_LIMIT = 2
MIMO_VALIDATION_RETRY_LIMIT = 1
ACTIVE_MODULE = "回顾性队列研究方案"
MODULE_OPTIONS = [
    ACTIVE_MODULE,
    "待开发模块 1",
    "待开发模块 2",
    "待开发模块 3",
    "待开发模块 4",
    "待开发模块 5",
]

THINKING_PROFILES = {
    "deepseek": [
        {"thinking": {"type": "enabled"}, "reasoning_effort": "max"},
        {"reasoning_effort": "max"},
        {"thinking": {"type": "enabled"}},
    ],
    "kimi": [
        {"thinking": {"type": "disabled"}},
    ],
    "GLM": [
        {"thinking": {"type": "enabled"}, "reasoning_effort": "max"},
        {"thinking": {"type": "enabled"}},
        {"reasoning_effort": "max"},
    ],
    "mimo": [
        {"thinking": {"type": "enabled"}, "reasoning_effort": "max", "reasoning": {"effort": "max"}},
        {"reasoning": {"effort": "max"}},
        {"thinking": {"type": "enabled"}},
        {"reasoning_effort": "max"},
    ],
}

CONFIG_DIR = Path(os.environ.get("APPDATA", str(Path.home()))) / "ProtocolDesign"
CONFIG_FILE = CONFIG_DIR / "settings.json"
CRYPTPROTECT_UI_FORBIDDEN = 0x1


class DATA_BLOB(ctypes.Structure):
    _fields_ = [("cbData", wintypes.DWORD), ("pbData", ctypes.POINTER(ctypes.c_byte))]


class KimiStreamNoContentError(RuntimeError):
    """Raised when Kimi stream returns reasoning events but no final answer content."""

    def __init__(self, reasoning_seen: bool, finish_reasons: list[str], usage: dict | None, chunks_seen: int) -> None:
        self.reasoning_seen = reasoning_seen
        self.finish_reasons = finish_reasons
        self.usage = usage
        self.chunks_seen = chunks_seen
        finish_text = "、".join(finish_reasons) if finish_reasons else "未返回"
        usage_text = json.dumps(usage, ensure_ascii=False) if usage else "未返回"
        super().__init__(
            "Kimi 流式响应未返回最终方案正文 content。"
            f"reasoning_seen={reasoning_seen}，finish_reason={finish_text}，usage={usage_text}，chunks={chunks_seen}。"
            "模型可能只返回了思考内容或在进入最终回答前中断。"
        )


@dataclass(frozen=True)
class GenerationWorkflow:
    module: str
    provider: str
    base_url: str
    api_key: str
    model: str
    input_path: Path
    output_dir: Path
    output_name: str
    remember_key: bool


crypt32 = ctypes.windll.crypt32
kernel32 = ctypes.windll.kernel32
crypt32.CryptProtectData.argtypes = [
    ctypes.POINTER(DATA_BLOB),
    wintypes.LPCWSTR,
    ctypes.POINTER(DATA_BLOB),
    wintypes.LPVOID,
    wintypes.LPVOID,
    wintypes.DWORD,
    ctypes.POINTER(DATA_BLOB),
]
crypt32.CryptProtectData.restype = wintypes.BOOL
crypt32.CryptUnprotectData.argtypes = [
    ctypes.POINTER(DATA_BLOB),
    ctypes.POINTER(wintypes.LPWSTR),
    ctypes.POINTER(DATA_BLOB),
    wintypes.LPVOID,
    wintypes.LPVOID,
    wintypes.DWORD,
    ctypes.POINTER(DATA_BLOB),
]
crypt32.CryptUnprotectData.restype = wintypes.BOOL
kernel32.LocalFree.argtypes = [wintypes.HLOCAL]
kernel32.LocalFree.restype = wintypes.HLOCAL

REQUIRED_SUMMARY_ROWS = [
    "题目",
    "研究中心",
    "研究目的",
    "研究设计",
    "研究人群",
    "暴露与对照",
    "结局指标",
    "样本量/可行性",
    "主要协变量",
    "统计分析",
    "伦理与数据保护",
]

REQUIRED_SECTION_HEADINGS = [
    "1. 研究背景与意义",
    "2. 研究目的",
    "2.1 主要研究目的",
    "2.2 次要研究目的",
    "3. 研究设计",
    "3.1 研究类型与总体框架",
    "3.2 设计选择的合理性",
    "3.3 时间零点（time zero）、基线期与随访",
    "3.4 偏倚最小化的设计考虑",
    "4. 研究对象",
    "4.1 数据来源",
    "4.2 目标人群、来源人群与队列构建",
    "4.3 纳入标准",
    "4.4 排除标准",
    "4.5 样本量估算",
    "5. 暴露因素定义与选择",
    "5.1 暴露定义的依据",
    "5.2 暴露的识别定义",
    "6. 对照组定义与选择",
    "6.1 对照组定义的依据",
    "6.2 对照组的识别定义",
    "7. 适应证、禁忌证和治疗强度的可比性",
    "8. 联合用药、换药、停药、依从性与误分类处理",
    "9. 结局指标定义与测量",
    "9.1 主要结局",
    "9.2 次要结局与安全性结局",
    "9.3 结局识别来源与验证",
    "9.4 重复事件与竞争事件",
    "9.5 替代定义、敏感性分析与特殊结局说明",
    "10. 协变量选择与测量",
    "10.1 协变量选择原则与因果框架",
    "10.2 人口学和社会经济学变量",
    "10.3 临床特征、疾病严重程度与合并症",
    "10.4 合并用药、合并治疗与医疗利用指标",
    "10.5 效应修饰因素与预设亚组变量",
    "10.6 协变量测量窗口与编码",
    "11. 统计分析方法",
    "11.1 总体分析原则与分析集",
    "11.2 描述性统计与基线平衡评估",
    "11.3 混杂控制策略与模型选择",
    "11.4 主要结局分析",
    "11.5 次要结局、亚组分析与效应异质性",
    "11.6 时间变化暴露、停药与依从性分析",
    "11.7 缺失数据处理",
    "11.8 敏感性分析",
    "11.9 统计软件、显著性水平与结果报告",
    "12. 数据管理",
    "12.1 数据提取、链接与去标识化",
    "12.2 数据管理与质量控制",
    "13. 伦理考量",
    "13.1 伦理审查",
    "13.2 知情同意获取方式或豁免",
    "13.3 隐私保护与去标识化",
    "13.4 利益冲突声明",
    "附录：建议附件与表单",
    "附表 1 暴露、对照、结局与协变量操作性定义总表",
    "附表 2 敏感性分析矩阵",
]

KIMI_SECTION_BATCHES = [
    REQUIRED_SECTION_HEADINGS[0:8],
    REQUIRED_SECTION_HEADINGS[8:15],
    REQUIRED_SECTION_HEADINGS[15:23],
    REQUIRED_SECTION_HEADINGS[23:31],
    REQUIRED_SECTION_HEADINGS[31:38],
    REQUIRED_SECTION_HEADINGS[38:45],
    REQUIRED_SECTION_HEADINGS[45:50],
    REQUIRED_SECTION_HEADINGS[50:],
]


def app_root() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parents[1]


def resource_path(*parts: str) -> Path:
    base = Path(getattr(sys, "_MEIPASS", app_root()))
    return base.joinpath(*parts)


def protect_text(value: str) -> str:
    data = value.encode("utf-8")
    in_buffer = ctypes.create_string_buffer(data)
    in_blob = DATA_BLOB(len(data), ctypes.cast(in_buffer, ctypes.POINTER(ctypes.c_byte)))
    out_blob = DATA_BLOB()
    if not crypt32.CryptProtectData(
        ctypes.byref(in_blob),
        None,
        None,
        None,
        None,
        CRYPTPROTECT_UI_FORBIDDEN,
        ctypes.byref(out_blob),
    ):
        raise ctypes.WinError()
    try:
        encrypted = ctypes.string_at(out_blob.pbData, out_blob.cbData)
        return base64.b64encode(encrypted).decode("ascii")
    finally:
        kernel32.LocalFree(ctypes.cast(out_blob.pbData, wintypes.HLOCAL))


def unprotect_text(value: str) -> str:
    encrypted = base64.b64decode(value.encode("ascii"))
    in_buffer = ctypes.create_string_buffer(encrypted)
    in_blob = DATA_BLOB(len(encrypted), ctypes.cast(in_buffer, ctypes.POINTER(ctypes.c_byte)))
    out_blob = DATA_BLOB()
    if not crypt32.CryptUnprotectData(
        ctypes.byref(in_blob),
        None,
        None,
        None,
        None,
        CRYPTPROTECT_UI_FORBIDDEN,
        ctypes.byref(out_blob),
    ):
        raise ctypes.WinError()
    try:
        data = ctypes.string_at(out_blob.pbData, out_blob.cbData)
        return data.decode("utf-8")
    finally:
        kernel32.LocalFree(ctypes.cast(out_blob.pbData, wintypes.HLOCAL))


def load_settings() -> dict:
    if not CONFIG_FILE.exists():
        return {"api_keys": {}}
    try:
        data = json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {"api_keys": {}}
    if not isinstance(data, dict):
        return {"api_keys": {}}
    if not isinstance(data.get("api_keys"), dict):
        data["api_keys"] = {}
    return data


def save_settings(data: dict) -> None:
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    CONFIG_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    try:
        os.chmod(CONFIG_FILE, 0o600)
    except OSError:
        pass


def saved_api_key(provider: str) -> str:
    data = load_settings()
    encrypted = data.get("api_keys", {}).get(provider)
    if not encrypted:
        return ""
    try:
        return unprotect_text(str(encrypted))
    except Exception:
        return ""


def has_saved_api_key(provider: str) -> bool:
    return bool(load_settings().get("api_keys", {}).get(provider))


def save_api_key(provider: str, api_key: str) -> None:
    data = load_settings()
    data.setdefault("api_keys", {})[provider] = protect_text(api_key)
    save_settings(data)


def delete_saved_api_key(provider: str) -> None:
    data = load_settings()
    if data.get("api_keys", {}).pop(provider, None) is not None:
        save_settings(data)


def read_text_file(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="ignore")


def read_docx(path: Path) -> str:
    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"\n[表格 {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", "；") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def read_summary(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return read_docx(path)
    if suffix in {".md", ".markdown", ".txt"}:
        return read_text_file(path)
    raise ValueError("暂仅支持 docx、md、txt 格式的方案摘要文档。")


def normalize_endpoint(base_url: str) -> str:
    cleaned = base_url.strip().rstrip("/")
    if not cleaned:
        raise ValueError("请填写 OpenAI 兼容接口的 base_url。")
    if cleaned.endswith("/chat/completions"):
        return cleaned
    return f"{cleaned}/chat/completions"




def build_prompt(summary_text: str) -> str:
    skill = resource_crypto.decrypt_resource_text("SKILL.md")
    guide = resource_crypto.decrypt_resource_text("references/template-writing-guide.md")
    headings = "\n".join(f"- {heading}" for heading in REQUIRED_SECTION_HEADINGS)
    summary_rows = "、".join(REQUIRED_SUMMARY_ROWS)
    return f"""
你是一名熟悉回顾性队列研究、真实世界研究、临床流行病学、统计分析计划和医学伦理的中文研究方案写作专家。
请严格按照下方 cohort-protocol 技能规则与模板写作指南，把用户提供的"方案摘要"扩写为完整、出版级、可供伦理审查和统计实施参考的中文回顾性队列研究方案。

输出要求：
1. 只输出一个合法 JSON 对象，不要 Markdown，不要代码块，不要解释文字。
2. JSON 必须符合 renderer schema：cover、summary、sections 三个顶层字段。
3. summary 必须是二维数组，行标签依次覆盖：{summary_rows}。
4. sections 必须按以下标题顺序生成，level=1 用于整数编号大节、附录和附表标题，level=2 用于二级标题：
{headings}
5. 每个主要方法学章节必须写出实质内容，不得只给提纲；缺失但关键的信息可作方法学上稳妥的默认推断，但不得出现占位符，也不得把不确定性直接写入完整版方案正文。
6. 对暴露、对照、time zero、基线期、随访、结局、协变量、混杂控制、缺失数据、敏感性分析、数据管理和伦理合规必须明确说明。
6a. 完整版方案正文中的核心研究要素必须使用明确、确定、可执行的表述，不得使用"约""大概""预计约""可能为"等模糊或不确定措辞；如用户摘要缺少关键细节，应采用方法学上稳妥且明确的默认设定，而不是暴露不确定表达。
6b. "3.3 时间零点（time zero）、基线期与随访"章节必须明确写出：随访起始时间、随访结束时间或截止日期、总随访时长（如"自索引日期起随访3年"或"观察索引日期后6年内结局指标的发生情况"）、以及删失规则；不得把总随访时长写成取决于纳入日期的范围表达，不得出现类似"总随访时间范围为0至6年（取决于纳入日期）"或"0–6年不等"的表述；不得使用"例如""比如""举例来说"等引导词，通过列举不同纳入日期/入组日期/索引日期患者的个体化例子来解释随访期，正文应直接用统一适用于全队列的规则句描述随访起止、固定观察窗和删失规则。
6c. "9. 结局指标定义与测量"及其子章节（9.1–9.5）必须为每个结局指标明确写出结局观察时间窗，即自索引日期（index date）起在明确的年数、月数或日数内识别该结局（如"自索引日期起3年内首次心衰住院"或"自索引日期起6个月内首次严重感染"），不得仅写"随访期间发生"而不限定观察窗长度；每个主要、次要和安全性结局都必须单独写清观察时间长度，不能只在总述中笼统带过。
7. 附表 1 使用 headers=["类别","变量名称","操作性定义","数据来源","时间窗","编码/算法","备注"]，rows 至少覆盖暴露、对照、主要结局、次要/安全性结局、核心协变量。
8. 附表 2 使用 headers=["敏感性分析编号","目的","变更内容","对应偏倚/假设","结果指标"]，rows 至少 5 行。
9. 禁止出现模板说明、写作要点、示例标签、【】占位符、方括号占位符和"若适用"等未清理痕迹。
10. 正文不需要加入软件风险提示，风险提示由软件界面和帮助文档承担。
11. 请参照内置方案模板示例的写作风格和写作要点，将方案写成可供伦理审查、统计实施和数据提取参考的详尽版本。整体中文字符数建议不少于 {TARGET_PROTOCOL_CHINESE_CHARS} 个，但不得机械重复或堆砌空话；应通过充分展开研究背景、操作性定义、统计分析、数据管理、伦理合规、质量控制和敏感性分析来提高完整性。
12. 本次调用已由软件默认启用大模型思考/推理模式，并请求最大思考强度；请在生成前进行充分方法学推理，但最终只输出 JSON，不输出思考过程。

技能说明：
{skill}

模板写作指南：
{guide}

用户方案摘要：
{summary_text}
""".strip()


def build_validation_retry_prompt(original_prompt: str, validation_error: str) -> str:
    headings = "\n".join(f"- {heading}" for heading in REQUIRED_SECTION_HEADINGS)
    return f"""
{original_prompt}

重要补充要求：
上一次模型输出未通过软件结构校验，失败原因是：{validation_error}
请重新生成一个完整 JSON 对象，必须保留 cover、summary、sections 三个顶层字段。
sections 不能只输出少数章节，必须完整覆盖以下全部章节标题，并按顺序输出：
{headings}
每个 section 应包含 heading、level、paragraphs；附表章节还应包含 tables。请不要省略后续章节，不要只输出摘要或提纲。
如果失败原因包含 Unterminated string、Invalid control character、JSONDecodeError 或类似 JSON 语法错误，说明上一次输出很可能被截断或字符串未正确转义；这次必须保证所有字符串、数组和对象完整闭合，段落内部换行必须写成 \\n，不允许输出裸换行、裸制表符或半截句子。必要时可以适度压缩单段文字，但必须优先保证 JSON 完整、章节完整、附表完整。
最终仍然只能输出严格合法 JSON 对象，不要 Markdown、代码块、解释文字或思考过程。
""".strip()


def infer_section_level(heading: str) -> int:
    if heading.startswith("附录") or heading.startswith("附表"):
        return 1
    if re.match(r"^\d+\.\d+", heading.strip()):
        return 2
    return 1


def build_kimi_cover_summary_prompt(summary_text: str) -> str:
    skill = resource_crypto.decrypt_resource_text("SKILL.md")
    guide = resource_crypto.decrypt_resource_text("references/template-writing-guide.md")
    summary_rows = "、".join(REQUIRED_SUMMARY_ROWS)
    return f"""
你是一名熟悉回顾性队列研究、真实世界研究、临床流行病学、统计分析计划和医学伦理的中文研究方案写作专家。
本次只生成完整方案 JSON 的 cover 和 summary 片段，软件会在后续步骤分批生成 sections 并本地组装。

输出要求：
1. 只输出一个合法 JSON 对象，不要 Markdown，不要代码块，不要解释文字，不要思考过程。
2. 顶层字段只能包含 cover 和 summary。
3. cover 必须是对象，至少包含：研究题目、研究类型、版本日期、研究中心、主要研究者、方案版本。
4. summary 必须是二维数组，行标签依次覆盖：{summary_rows}。其中"结局指标"行应简要概括主要结局、次要结局及其观察时间窗（如"自索引日期起X年内首次心衰住院；自索引日期起X年内全因死亡"），并使用明确表述，不得出现"约""大概"等模糊用词；不得仅罗列结局名称而不写时间窗，也不得把观察期写成随纳入时间变化的范围。
5. 内容应参照方案模板示例的写作风格，写成可供伦理审查、统计实施和数据提取参考的正式表达。
6. 核心研究要素应使用明确、确定的表述；若用户摘要缺少细节，应采用方法学上稳妥的明确默认设定，不得直接写入不确定表达。
7. 段落内部换行必须写成 \\n，不允许输出裸换行、裸制表符或半截句子。

技能说明：
{skill}

模板写作指南：
{guide}

用户方案摘要：
{summary_text}
""".strip()


def build_kimi_sections_prompt(
    summary_text: str,
    headings_batch: list[str],
    cover: dict,
    summary_rows: list[list[str]],
    generated_headings: list[str],
) -> str:
    skill = resource_crypto.decrypt_resource_text("SKILL.md")
    guide = resource_crypto.decrypt_resource_text("references/template-writing-guide.md")
    requested_headings = json.dumps(headings_batch, ensure_ascii=False)
    requested_levels = json.dumps(
        {heading: infer_section_level(heading) for heading in headings_batch},
        ensure_ascii=False,
    )
    prior = "、".join(generated_headings) if generated_headings else "无"
    cover_text = json.dumps(cover, ensure_ascii=False)
    summary_text_json = json.dumps(summary_rows, ensure_ascii=False)
    return f"""
你是一名熟悉回顾性队列研究、真实世界研究、临床流行病学、统计分析计划和医学伦理的中文研究方案写作专家。
本次只生成完整方案 JSON 的 sections 片段，软件会把多个片段本地组装成最终方案。

输出要求：
1. 只输出一个合法 JSON 对象，不要 Markdown，不要代码块，不要解释文字，不要思考过程。
2. 顶层字段只能包含 sections。
3. sections 必须只包含"本批次唯一允许章节标题 JSON 数组"中的章节，并严格按数组顺序返回，不得缺失、合并、改名或额外增加章节。
4. 每个 section 必须包含 heading、level、paragraphs；heading 必须与请求标题完全一致；level 必须使用"本批次章节 level 映射"中指定的数值。
5. paragraphs 必须是字符串数组。每个主要方法学章节必须写出实质内容，不得只给提纲；缺失但关键的信息可作方法学上稳妥的默认推断，但不得出现占位符，也不得把不确定性直接写入完整版方案正文。
6. 对暴露、对照、time zero、基线期、随访、结局、协变量、混杂控制、缺失数据、敏感性分析、数据管理和伦理合规等相关章节必须明确说明。
6a. 完整版方案正文中的核心研究要素必须使用明确、确定、可执行的表述，不得使用"约""大概""预计约""可能为"等模糊或不确定措辞；如用户摘要缺少关键细节，应采用方法学上稳妥且明确的默认设定，而不是暴露不确定表达。
6b. 撰写"3.3 时间零点（time zero）、基线期与随访"时，必须明确写出：随访起始时间、随访结束时间或截止日期、总随访时长（如"自索引日期起随访3年"或"观察索引日期后6年内结局指标的发生情况"）、以及删失规则；不得仅写"随访至结局发生或数据截止"等模糊表述，也不得写成类似"总随访时间范围为0至6年（取决于纳入日期）"、"0–6年不等"这类随纳入时间变化的范围表达；不得使用"例如""比如""举例来说"等引导词，通过列举不同纳入日期/入组日期/索引日期患者的个体化例子来解释随访期，正文应直接用统一适用于全队列的规则句描述随访起止、固定观察窗和删失规则。
6c. 撰写"9. 结局指标定义与测量"及其子章节（9.1–9.5）时，必须为每个结局指标明确写出结局观察时间窗，即自索引日期（index date）起在明确的年数、月数或日数内识别该结局（如"自索引日期起3年内首次心衰住院"或"自索引日期起30日内首次再入院"），不得仅写"随访期间发生"而不限定观察窗长度；每个主要、次要和安全性结局都必须单独写清观察时间长度，不能只在总述中笼统带过。
7. 附表 1 必须包含 tables，且使用 headers=["类别","变量名称","操作性定义","数据来源","时间窗","编码/算法","备注"]，rows 至少覆盖暴露、对照、主要结局、次要/安全性结局、核心协变量。
8. 附表 2 必须包含 tables，且使用 headers=["敏感性分析编号","目的","变更内容","对应偏倚/假设","结果指标"]，rows 至少 5 行。
9. 禁止出现模板说明、写作要点、示例标签、【】占位符、方括号占位符和"若适用"等未清理痕迹。
10. 请参照内置方案模板示例的写作风格和写作要点，充分展开本批次章节；但必须优先保证 JSON 完整闭合。
11. 段落内部换行必须写成 \\n，不允许输出裸换行、裸制表符或半截句子。

本批次唯一允许章节标题 JSON 数组：
{requested_headings}

本批次章节 level 映射：
{requested_levels}

已生成章节：{prior}

方案封面上下文：
{cover_text}

方案摘要上下文：
{summary_text_json}

技能说明：
{skill}

模板写作指南：
{guide}

用户方案摘要：
{summary_text}
""".strip()


def build_kimi_segment_retry_prompt(segment_prompt: str, stage_name: str, validation_error: str) -> str:
    return f"""
{segment_prompt}

重要补充要求：
上一次 Kimi 分段输出未通过软件校验，失败阶段是：{stage_name}
失败原因是：{validation_error}
请只重新生成当前阶段要求的严格合法 JSON，不要重写其他阶段。
如果失败原因包含 Unterminated string、JSONDecodeError、Invalid control character 或类似 JSON 语法错误，说明上一次片段可能被截断或字符串未正确转义；这次必须保证所有字符串、数组和对象完整闭合。
最终仍然只能输出一个 JSON 对象，不要 Markdown、代码块、解释文字或思考过程。
""".strip()


def should_continue_response_format_fallback(message: str) -> bool:
    lowered = message.lower()
    return "response_format" in lowered or "400" in message


def should_continue_compatibility_payload_fallback(message: str) -> bool:
    if should_continue_response_format_fallback(message):
        return True
    lowered = message.lower()
    compatibility_markers = (
        "模型接口返回错误 415",
        "模型接口返回错误 422",
        "unsupported",
        "not support",
        "invalid parameter",
        "invalid_request",
        "unrecognized",
        "不支持",
        "参数",
        "thinking",
        "reasoning",
    )
    return any(marker in lowered or marker in message for marker in compatibility_markers)


def should_continue_kimi_fallback(message: str) -> bool:
    if should_continue_response_format_fallback(message):
        return True
    lowered = message.lower()
    retryable_markers = (
        "模型接口连接失败",
        "模型流式响应解析失败",
        "模型接口返回的 json 无法解析",
        "模型流式返回中未找到最终内容",
        "未返回最终方案正文",
        "message.content",
        "engine_overloaded",
        "overloaded",
        "too many requests",
        "rate limit",
        "temporarily unavailable",
        "timeout",
        "请求过于频繁",
        "稍后",
        "限制",
    )
    return any(marker in lowered for marker in retryable_markers)


def reasoning_attempts(profiles: list[dict], temperature: float = 0.2, max_tokens: int | None = None) -> list[dict]:
    attempts: list[dict] = []
    for options in profiles:
        attempts.append(
            {
                "label": "JSON 模式",
                "stream": False,
                "use_response_format": True,
                "thinking_options": options,
                "temperature": temperature,
                "max_tokens": max_tokens,
            }
        )
        attempts.append(
            {
                "label": "普通模式",
                "stream": False,
                "use_response_format": False,
                "thinking_options": options,
                "temperature": temperature,
                "max_tokens": max_tokens,
            }
        )
    return attempts


def is_transient_model_error(status_code: int, response_text: str) -> bool:
    if status_code == 429:
        return True
    if status_code not in TRANSIENT_STATUS_CODES:
        return False
    lowered = response_text.lower()
    transient_markers = [
        "engine_overloaded",
        "overloaded",
        "try again later",
        "rate limit",
        "too many requests",
        "temporarily unavailable",
        "timeout",
        "1302",
        "频率",
        "限制",
        "稍后",
        "请求过于频繁",
        "调用限制",
    ]
    return any(marker in lowered for marker in transient_markers)


def is_transient_transport_error(exc: requests.exceptions.RequestException) -> bool:
    return isinstance(exc, (requests.exceptions.SSLError, requests.exceptions.ConnectionError, requests.exceptions.Timeout))


def decode_stream_line(raw_line) -> str:
    if isinstance(raw_line, bytes):
        try:
            return raw_line.decode("utf-8")
        except UnicodeDecodeError:
            return raw_line.decode("utf-8", errors="replace")
    return str(raw_line)


def looks_like_mojibake(text: str) -> bool:
    if not text:
        return False
    markers = ("ä", "å", "ç", "è", "é", "æ", "ï¼", "ã€", "â", "�")
    marker_count = sum(text.count(marker) for marker in markers)
    chinese_count = len(re.findall(r"[\u4e00-\u9fff]", text))
    return marker_count >= 6 and chinese_count < marker_count


def repair_mojibake_text(text: str) -> str:
    if not looks_like_mojibake(text):
        return text
    for source_encoding in ("latin1", "cp1252"):
        try:
            repaired = text.encode(source_encoding).decode("utf-8")
        except UnicodeError:
            continue
        if len(re.findall(r"[\u4e00-\u9fff]", repaired)) > len(re.findall(r"[\u4e00-\u9fff]", text)):
            return repaired
    return text


def extract_content_from_stream(response: requests.Response) -> str:
    parts: list[str] = []
    reasoning_seen = False
    finish_reasons: list[str] = []
    usage: dict | None = None
    chunks_seen = 0
    response.encoding = "utf-8"
    for raw_line in response.iter_lines(decode_unicode=False):
        if not raw_line:
            continue
        line = decode_stream_line(raw_line).strip()
        if line.startswith("data:"):
            line = line[5:].strip()
        if not line or line == "[DONE]":
            continue
        try:
            chunk = json.loads(line)
        except json.JSONDecodeError:
            continue
        chunks_seen += 1
        if isinstance(chunk.get("usage"), dict):
            usage = chunk["usage"]
        for choice in chunk.get("choices", []):
            finish_reason = choice.get("finish_reason")
            if finish_reason:
                finish_reasons.append(str(finish_reason))
            delta = choice.get("delta") or {}
            if delta.get("reasoning_content"):
                reasoning_seen = True
            content = delta.get("content")
            if content:
                parts.append(content)
            message = choice.get("message") or {}
            if message.get("reasoning_content"):
                reasoning_seen = True
            message_content = message.get("content")
            if message_content:
                parts.append(message_content)
    content = repair_mojibake_text("".join(parts).strip())
    if not content:
        if reasoning_seen:
            raise KimiStreamNoContentError(reasoning_seen, finish_reasons, usage, chunks_seen)
        raise RuntimeError("模型流式返回中未找到最终内容。")
    return content


def build_payload(
    model: str,
    prompt: str,
    use_response_format: bool,
    stream: bool,
    thinking_options: dict,
    temperature: float,
    max_tokens: int | None = None,
) -> dict:
    payload = {
        "model": model.strip(),
        "messages": [
            {
                "role": "system",
                "content": "你只输出严格合法的 JSON 对象。不要输出 Markdown、代码块、思考过程或额外解释。",
            },
            {"role": "user", "content": prompt},
        ],
        "temperature": temperature,
        "stream": stream,
    }
    if max_tokens is not None:
        payload["max_tokens"] = max_tokens
    payload.update(thinking_options)
    if use_response_format:
        payload["response_format"] = {"type": "json_object"}
    return payload


def post_chat_completion(
    base_url: str,
    api_key: str,
    model: str,
    prompt: str,
    use_response_format: bool,
    thinking_options: dict,
    stream: bool = False,
    temperature: float = 0.2,
    max_tokens: int | None = None,
    reasoning_only_means_no_content: bool = False,
) -> str:
    endpoint = normalize_endpoint(base_url)
    headers = {
        "Authorization": f"Bearer {api_key.strip()}",
        "Content-Type": "application/json",
        "Accept": "application/json, text/event-stream; charset=utf-8",
        "Accept-Charset": "utf-8",
    }
    payload = build_payload(model, prompt, use_response_format, stream, thinking_options, temperature, max_tokens)
    response = None
    last_transport_error: requests.exceptions.RequestException | None = None
    for attempt in range(len(TRANSIENT_RETRY_DELAYS) + 1):
        try:
            response = requests.post(
                endpoint,
                headers=headers,
                json=payload,
                timeout=MODEL_REQUEST_TIMEOUT_SECONDS,
                stream=payload["stream"],
            )
        except requests.exceptions.RequestException as exc:
            last_transport_error = exc
            if attempt >= len(TRANSIENT_RETRY_DELAYS) or not is_transient_transport_error(exc):
                raise RuntimeError(f"模型接口连接失败：{exc}") from exc
            time.sleep(TRANSIENT_RETRY_DELAYS[attempt])
            continue
        if response.status_code < 400:
            last_transport_error = None
            break
        response_text = response.text[:1200]
        if attempt >= len(TRANSIENT_RETRY_DELAYS) or not is_transient_model_error(response.status_code, response_text):
            break
        response.close()
        time.sleep(TRANSIENT_RETRY_DELAYS[attempt])
    if response is None:
        if last_transport_error is not None:
            raise RuntimeError(f"模型接口连接失败：{last_transport_error}") from last_transport_error
        raise RuntimeError("模型接口未返回响应。")
    try:
        if response.status_code >= 400:
            raise RuntimeError(f"模型接口返回错误 {response.status_code}：{response.text[:1200]}")
        if payload["stream"]:
            try:
                return extract_content_from_stream(response)
            except RuntimeError:
                raise
            except (requests.exceptions.RequestException, ValueError) as exc:
                raise RuntimeError(f"模型流式响应解析失败：{exc}") from exc
        try:
            data = response.json()
        except (requests.exceptions.RequestException, json.JSONDecodeError, ValueError) as exc:
            raise RuntimeError(f"模型接口返回的 JSON 无法解析：{exc}") from exc
        try:
            choice = data["choices"][0]
            message = choice["message"]
            content = message.get("content")
            if content:
                return content
            if reasoning_only_means_no_content and message.get("reasoning_content"):
                finish_reason = choice.get("finish_reason")
                finish_reasons = [str(finish_reason)] if finish_reason else []
                usage = data.get("usage") if isinstance(data.get("usage"), dict) else None
                raise KimiStreamNoContentError(True, finish_reasons, usage, 1)
            raise RuntimeError(f"模型接口返回中未找到 message.content：{data}")
        except (KeyError, IndexError, TypeError) as exc:
            raise RuntimeError(f"模型接口返回结构不符合 OpenAI chat/completions 格式：{data}") from exc
    finally:
        response.close()


def call_standard_model_workflow(
    base_url: str,
    api_key: str,
    model: str,
    prompt: str,
    attempts: list[dict],
    should_continue,
    failure_message: str,
) -> str:
    errors: list[str] = []
    for attempt in attempts:
        try:
            return post_chat_completion(
                base_url,
                api_key,
                model,
                prompt,
                attempt["use_response_format"],
                attempt["thinking_options"],
                stream=attempt["stream"],
                temperature=attempt["temperature"],
                max_tokens=attempt["max_tokens"],
                reasoning_only_means_no_content=attempt.get("reasoning_only_means_no_content", False),
            )
        except RuntimeError as exc:
            message = str(exc)
            errors.append(f"{attempt['label']}：{message}")
            if not should_continue(message):
                raise
    detail = "\n".join(f"- {item}" for item in errors[-4:])
    raise RuntimeError(f"{failure_message}\n{detail}")


def call_deepseek_model(base_url: str, api_key: str, model: str, prompt: str) -> str:
    attempts = reasoning_attempts(THINKING_PROFILES["deepseek"])
    return call_standard_model_workflow(
        base_url,
        api_key,
        model,
        prompt,
        attempts,
        should_continue_response_format_fallback,
        "DeepSeek 模型接口未接受当前思考模式/最大 effort 请求，请核对模型是否支持 thinking/reasoning 参数。",
    )


def call_glm_model(base_url: str, api_key: str, model: str, prompt: str) -> str:
    attempts = reasoning_attempts(THINKING_PROFILES["GLM"])
    return call_standard_model_workflow(
        base_url,
        api_key,
        model,
        prompt,
        attempts,
        should_continue_response_format_fallback,
        "GLM 模型接口未接受当前思考模式/最大 effort 请求，请核对模型是否支持 thinking/reasoning 参数。",
    )


def call_mimo_model(base_url: str, api_key: str, model: str, prompt: str) -> str:
    attempts = reasoning_attempts(THINKING_PROFILES["mimo"])
    return call_standard_model_workflow(
        base_url,
        api_key,
        model,
        prompt,
        attempts,
        should_continue_response_format_fallback,
        "MiMo 模型接口未接受当前思考模式/最大 effort 请求，请核对模型是否支持 thinking/reasoning 参数。",
    )


def call_custom_model(base_url: str, api_key: str, model: str, prompt: str) -> str:
    attempts = [
        {
            "label": "普通 OpenAI 兼容模式",
            "stream": False,
            "use_response_format": False,
            "thinking_options": {},
            "temperature": 0.2,
            "max_tokens": None,
        },
        {
            "label": "普通 OpenAI 兼容 JSON 模式",
            "stream": False,
            "use_response_format": True,
            "thinking_options": {},
            "temperature": 0.2,
            "max_tokens": None,
        },
    ]
    return call_standard_model_workflow(
        base_url,
        api_key,
        model,
        prompt,
        attempts,
        should_continue_compatibility_payload_fallback,
        "模型接口未接受当前请求参数；已按最大兼容性尝试普通 OpenAI 兼容负载和 JSON 模式。",
    )


def call_model(base_url: str, api_key: str, model: str, prompt: str, provider: str) -> str:
    if not api_key.strip():
        raise ValueError("请填写 api_key。")
    if not model.strip():
        raise ValueError("请填写模型名称。")
    if provider == "deepseek":
        return call_deepseek_model(base_url, api_key, model, prompt)
    if provider == "kimi":
        return call_kimi_model(base_url, api_key, model, prompt)
    if provider == "GLM":
        return call_glm_model(base_url, api_key, model, prompt)
    if provider == "mimo":
        return call_mimo_model(base_url, api_key, model, prompt)
    return call_custom_model(base_url, api_key, model, prompt)


def call_kimi_model(base_url: str, api_key: str, model: str, prompt: str) -> str:
    attempts = [
        {
            "stream": True,
            "use_response_format": False,
            "thinking_options": {"thinking": {"type": "disabled"}},
            "label": "流式禁用思考",
        },
        {
            "stream": True,
            "use_response_format": True,
            "thinking_options": {"thinking": {"type": "disabled"}},
            "label": "流式禁用思考 JSON 模式",
        },
        {
            "stream": False,
            "use_response_format": False,
            "thinking_options": {"thinking": {"type": "disabled"}},
            "label": "非流式禁用思考",
        },
    ]
    errors: list[str] = []
    for attempt in attempts:
        try:
            return post_chat_completion(
                base_url,
                api_key,
                model,
                prompt,
                attempt["use_response_format"],
                attempt["thinking_options"],
                stream=attempt["stream"],
                temperature=0.6,
                max_tokens=32768,
                reasoning_only_means_no_content=True,
            )
        except KimiStreamNoContentError as exc:
            errors.append(f"{attempt['label']}：{exc}")
            continue
        except RuntimeError as exc:
            message = str(exc)
            errors.append(f"{attempt['label']}：{message}")
            if should_continue_kimi_fallback(message):
                continue
            raise
    detail = "\n".join(f"- {item}" for item in errors[-4:])
    if any("未返回最终方案正文" in item or "reasoning_seen=True" in item for item in errors):
        raise RuntimeError(
            "Kimi 模型未返回最终方案正文；已按禁用思考模式尝试流式、JSON 模式和非流式调用。"
            "请稍后重试或切换模型。\n"
            f"{detail}"
        )
    raise RuntimeError(
        "Kimi 模型接口未接受当前请求参数；已按禁用思考模式尝试流式、JSON 模式和非流式调用。\n"
        f"{detail}"
    )


def call_kimi_fragment_model(base_url: str, api_key: str, model: str, prompt: str) -> str:
    attempts = [
        {
            "stream": True,
            "use_response_format": False,
            "thinking_options": {"thinking": {"type": "disabled"}},
            "label": "分段流式禁用思考",
        },
        {
            "stream": False,
            "use_response_format": False,
            "thinking_options": {"thinking": {"type": "disabled"}},
            "label": "分段非流式禁用思考",
        },
    ]
    errors: list[str] = []
    for attempt in attempts:
        try:
            return post_chat_completion(
                base_url,
                api_key,
                model,
                prompt,
                attempt["use_response_format"],
                attempt["thinking_options"],
                stream=attempt["stream"],
                temperature=0.6,
                max_tokens=32768,
                reasoning_only_means_no_content=True,
            )
        except KimiStreamNoContentError as exc:
            errors.append(f"{attempt['label']}：{exc}")
            continue
        except RuntimeError as exc:
            message = str(exc)
            errors.append(f"{attempt['label']}：{message}")
            if should_continue_kimi_fallback(message):
                continue
            raise
    detail = "\n".join(f"- {item}" for item in errors[-4:])
    raise RuntimeError(
        "Kimi 分段生成接口未能返回可用正文；已按禁用思考模式尝试流式和非流式调用。\n"
        f"{detail}"
    )


def extract_json_object(text: str) -> dict:
    cleaned = text.strip()
    cleaned = repair_mojibake_text(cleaned)
    fenced = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", cleaned, re.S)
    if fenced:
        cleaned = fenced.group(1).strip()
    if not cleaned.startswith("{"):
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start >= 0 and end > start:
            cleaned = cleaned[start : end + 1]
    data = loads_json_with_control_char_repair(cleaned)
    if not isinstance(data, dict):
        raise ValueError("模型输出不是 JSON 对象。")
    repaired_data = repair_mojibake_values(data)
    return repaired_data if isinstance(repaired_data, dict) else data


def repair_mojibake_values(value):
    if isinstance(value, str):
        return repair_mojibake_text(value)
    if isinstance(value, list):
        return [repair_mojibake_values(item) for item in value]
    if isinstance(value, dict):
        return {
            repair_mojibake_text(str(key)) if isinstance(key, str) else key: repair_mojibake_values(item)
            for key, item in value.items()
        }
    return value


def escape_control_chars_in_json_strings(text: str) -> str:
    parts: list[str] = []
    in_string = False
    escaped = False
    for char in text:
        if escaped:
            parts.append(char)
            escaped = False
            continue
        if char == "\\" and in_string:
            parts.append(char)
            escaped = True
            continue
        if char == '"':
            parts.append(char)
            in_string = not in_string
            continue
        if in_string and ord(char) < 0x20:
            if char == "\n":
                parts.append("\\n")
            elif char == "\r":
                parts.append("\\r")
            elif char == "\t":
                parts.append("\\t")
            else:
                parts.append(f"\\u{ord(char):04x}")
            continue
        parts.append(char)
    return "".join(parts)


def remove_trailing_commas_in_json(text: str) -> str:
    parts: list[str] = []
    in_string = False
    escaped = False
    length = len(text)
    for index, char in enumerate(text):
        if escaped:
            parts.append(char)
            escaped = False
            continue
        if char == "\\" and in_string:
            parts.append(char)
            escaped = True
            continue
        if char == '"':
            parts.append(char)
            in_string = not in_string
            continue
        if char == "," and not in_string:
            lookahead = index + 1
            while lookahead < length and text[lookahead] in " \t\r\n":
                lookahead += 1
            if lookahead < length and text[lookahead] in "]}":
                continue
        parts.append(char)
    return "".join(parts)


def strip_disallowed_control_chars(text: str) -> str:
    return "".join(
        char for char in text
        if ord(char) >= 0x20 or char in "\t\r\n"
    )


def loads_json_with_control_char_repair(text: str) -> dict:
    try:
        return json.loads(text)
    except json.JSONDecodeError as first_error:
        candidates = list(dict.fromkeys([
            escape_control_chars_in_json_strings(text),
            strip_disallowed_control_chars(text),
            remove_trailing_commas_in_json(text),
            strip_disallowed_control_chars(escape_control_chars_in_json_strings(text)),
            remove_trailing_commas_in_json(escape_control_chars_in_json_strings(text)),
            remove_trailing_commas_in_json(strip_disallowed_control_chars(text)),
            remove_trailing_commas_in_json(strip_disallowed_control_chars(escape_control_chars_in_json_strings(text))),
        ]))
        for repaired in candidates:
            if repaired == text:
                continue
            try:
                return json.loads(repaired)
            except json.JSONDecodeError:
                continue
        raise first_error


FIXED_TIME_WINDOW_PATTERN = re.compile(
    r"(?:自[^，。；\n]{0,20}?起\s*)?\d+(?:\.\d+)?\s*(?:个)?(?:年|月|周|天|日|小时)(?:内|以内|时|期间)?"
    r"|随访\s*\d+(?:\.\d+)?\s*(?:个)?(?:年|月|周|天|日|小时)"
    r"|观察[^，。；\n]{0,20}?\d+(?:\.\d+)?\s*(?:个)?(?:年|月|周|天|日|小时)内"
)
RANGE_TIME_WINDOW_PATTERN = re.compile(r"\d+(?:\.\d+)?\s*(?:-|–|—|~|至|到)\s*\d+(?:\.\d+)?\s*(?:个)?(?:年|月|周|天|日|小时)")
DISALLOWED_RANGE_FOLLOWUP_MARKERS = (
    "取决于纳入日期",
    "取决于入组日期",
    "随纳入时间变化",
    "随入组时间变化",
    "0至6年",
    "0-6年",
    "0–6年",
    "0—6年",
)
UNBOUNDED_OUTCOME_PHRASES = (
    "随访期间发生",
    "观察期间发生",
    "研究期间发生",
    "随访期内发生",
    "观察期内发生",
    "随访内发生",
)
FOLLOWUP_EXAMPLE_PATTERN = re.compile(
    r"(?:例如|比如|举例(?:来说)?)[^。；\n]{0,80}?对于[^。；\n]{0,80}?(?:纳入|入组|索引)[^。；\n]{0,80}?患者"
    r"|对于\s*\d{4}年\d{1,2}月\d{1,2}日[^。；\n]{0,80}?(?:纳入|入组|索引)[^。；\n]{0,80}?患者"
)


def normalize_text_block(value) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return "\n".join(str(item).strip() for item in value if str(item).strip())
    return str(value).strip()


def find_summary_row_text(summary, label: str) -> str:
    for row_label, row_value in normalize_summary_data(summary):
        if row_label == label:
            return row_value.strip()
    return ""


def find_section_text(sections, heading: str) -> str:
    if not isinstance(sections, list):
        return ""
    for section in sections:
        if not isinstance(section, dict):
            continue
        if str(section.get("heading", "")).strip() != heading:
            continue
        return normalize_text_block(section.get("paragraphs"))
    return ""


def contains_fixed_time_window(text: str) -> bool:
    normalized = text.replace("　", " ")
    return bool(FIXED_TIME_WINDOW_PATTERN.search(normalized))


def contains_disallowed_followup_range(text: str) -> bool:
    normalized = text.replace("　", " ")
    if any(marker in normalized for marker in DISALLOWED_RANGE_FOLLOWUP_MARKERS):
        return True
    return bool(RANGE_TIME_WINDOW_PATTERN.search(normalized) and any(marker in normalized for marker in ("范围", "不等", "取决于", "随纳入", "随入组")))


def contains_unbounded_outcome_phrase(text: str) -> bool:
    return any(phrase in text for phrase in UNBOUNDED_OUTCOME_PHRASES)


def contains_followup_example_phrase(text: str) -> bool:
    normalized = text.replace("　", " ")
    return bool(FOLLOWUP_EXAMPLE_PATTERN.search(normalized))


def validate_followup_and_outcome_windows(summary, sections) -> None:
    summary_outcomes = find_summary_row_text(summary, "结局指标")
    if not summary_outcomes:
        raise ValueError("summary 缺少“结局指标”行内容。")
    if contains_disallowed_followup_range(summary_outcomes):
        raise ValueError("summary 的“结局指标”行使用了随纳入时间变化的范围型观察期表达。")
    if not contains_fixed_time_window(summary_outcomes):
        raise ValueError("summary 的“结局指标”行必须为主要/次要结局写出明确观察时间窗（年/月/日）。")

    if not isinstance(sections, list):
        return

    followup_text = find_section_text(sections, "3.3 时间零点（time zero）、基线期与随访")
    if not followup_text:
        raise ValueError("缺少“3.3 时间零点（time zero）、基线期与随访”章节内容。")
    if contains_disallowed_followup_range(followup_text):
        raise ValueError("3.3 时间零点（time zero）、基线期与随访 不能写成随纳入时间变化的范围型随访时长。")
    if contains_followup_example_phrase(followup_text):
        raise ValueError("3.3 时间零点（time zero）、基线期与随访 不得通过列举不同纳入日期患者的例子解释随访规则，应改为统一适用于全队列的规则表述。")
    if not contains_fixed_time_window(followup_text):
        raise ValueError("3.3 时间零点（time zero）、基线期与随访 必须写出明确固定的随访时长或观察窗（年/月/日）。")

    main_outcome_text = find_section_text(sections, "9.1 主要结局")
    if not main_outcome_text:
        raise ValueError("缺少“9.1 主要结局”章节内容。")
    if contains_disallowed_followup_range(main_outcome_text):
        raise ValueError("9.1 主要结局 不得使用随纳入时间变化的范围型观察窗表达。")
    if contains_unbounded_outcome_phrase(main_outcome_text) and not contains_fixed_time_window(main_outcome_text):
        raise ValueError("9.1 主要结局 使用了“随访期间发生”类无界表述，但未写明明确观察时间窗。")
    if not contains_fixed_time_window(main_outcome_text):
        raise ValueError("9.1 主要结局 必须明确写出观察时间窗（年/月/日）。")

    secondary_outcome_text = find_section_text(sections, "9.2 次要结局与安全性结局")
    if not secondary_outcome_text:
        raise ValueError("缺少“9.2 次要结局与安全性结局”章节内容。")
    if contains_disallowed_followup_range(secondary_outcome_text):
        raise ValueError("9.2 次要结局与安全性结局 不得使用随纳入时间变化的范围型观察窗表达。")
    if contains_unbounded_outcome_phrase(secondary_outcome_text) and not contains_fixed_time_window(secondary_outcome_text):
        raise ValueError("9.2 次要结局与安全性结局 使用了“随访期间发生”类无界表述，但未写明明确观察时间窗。")
    if not contains_fixed_time_window(secondary_outcome_text):
        raise ValueError("9.2 次要结局与安全性结局 必须明确写出观察时间窗（年/月/日）。")



def collect_text_values(value) -> list[str]:
    if isinstance(value, str):
        return [value]
    if isinstance(value, list):
        texts: list[str] = []
        for item in value:
            texts.extend(collect_text_values(item))
        return texts
    if isinstance(value, dict):
        texts: list[str] = []
        for key, item in value.items():
            texts.extend(collect_text_values(key))
            texts.extend(collect_text_values(item))
        return texts
    return []


def collect_template_artifacts(value) -> list[str]:
    texts = collect_text_values(value)
    offenders: list[str] = []
    for pattern in protocol_renderer.FORBIDDEN_PATTERNS:
        if any(pattern in text for text in texts):
            offenders.append(pattern)
    bracket_placeholders: list[str] = []
    square_placeholders: list[str] = []
    for text in texts:
        bracket_placeholders.extend(re.findall(r"【[^】\r\n]{1,40}】", text))
        square_placeholders.extend(re.findall(r"\[[^\]\r\n]{0,40}[一-鿿][^\]\r\n]{0,40}\]", text))
    if bracket_placeholders:
        offenders.extend(sorted(set(bracket_placeholders))[:10])
    if square_placeholders:
        offenders.extend(sorted(set(square_placeholders))[:10])
    return offenders


def validate_protocol_json(data: dict) -> None:
    missing = [key for key in ("cover", "summary", "sections") if key not in data]
    if missing:
        raise ValueError("模型输出缺少字段：" + "、".join(missing))
    if not isinstance(data["sections"], list) or len(data["sections"]) < 20:
        raise ValueError("模型输出的 sections 数量不足，无法形成完整方案。")
    offenders = collect_template_artifacts(data)
    if offenders:
        raise ValueError("模型输出仍包含模板痕迹：" + "、".join(offenders))
    validate_followup_and_outcome_windows(data.get("summary"), data.get("sections"))


def model_output_excerpt(text: str, limit: int = 1000) -> str:
    cleaned = repair_mojibake_text(text.strip())
    if len(cleaned) <= limit * 2:
        return cleaned
    return f"{cleaned[:limit]}\n...\n{cleaned[-limit:]}"


def normalize_summary_data(summary) -> list[list[str]]:
    if isinstance(summary, dict):
        return [[str(key), str(value)] for key, value in summary.items()]
    if not isinstance(summary, list):
        raise ValueError("summary 必须是二维数组或对象。")
    rows: list[list[str]] = []
    for row in summary:
        if not isinstance(row, list) or len(row) < 2:
            raise ValueError("summary 必须是二维数组，每行至少包含标签和值。")
        rows.append([str(row[0]), str(row[1])])
    return rows


def validate_kimi_cover_summary_fragment(data: dict) -> tuple[dict, list[list[str]]]:
    cover = data.get("cover")
    summary = data.get("summary")
    if not isinstance(cover, dict) or not cover:
        raise ValueError("cover_summary 片段缺少有效 cover 对象。")
    summary_rows = normalize_summary_data(summary)
    labels = [row[0] for row in summary_rows]
    missing_rows = [row for row in REQUIRED_SUMMARY_ROWS if row not in labels]
    if missing_rows:
        raise ValueError("cover_summary 片段 summary 缺少行：" + "、".join(missing_rows))
    summary_outcomes = find_summary_row_text(summary_rows, "结局指标")
    if not summary_outcomes:
        raise ValueError("cover_summary 片段的“结局指标”行不能为空。")
    if contains_disallowed_followup_range(summary_outcomes):
        raise ValueError("cover_summary 片段的“结局指标”行不得使用随纳入时间变化的范围型观察期表达。")
    if not contains_fixed_time_window(summary_outcomes):
        raise ValueError("cover_summary 片段的“结局指标”行必须写出明确观察时间窗（年/月/日）。")
    return cover, summary_rows


def validate_table_payload(table, heading: str) -> None:
    if not isinstance(table, dict):
        raise ValueError(f"{heading} 的 tables 项必须是对象。")
    headers = table.get("headers")
    rows = table.get("rows")
    if not isinstance(headers, list) or not headers:
        raise ValueError(f"{heading} 的 table.headers 缺失或为空。")
    if not isinstance(rows, list) or not rows:
        raise ValueError(f"{heading} 的 table.rows 缺失或为空。")


def normalize_kimi_section(section: dict, expected_heading: str) -> dict:
    if not isinstance(section, dict):
        raise ValueError(f"{expected_heading} 对应 section 不是对象。")
    heading = str(section.get("heading", "")).strip()
    if heading != expected_heading:
        raise ValueError(f"章节标题不匹配：期望“{expected_heading}”，实际“{heading or '空'}”。")
    normalized = dict(section)
    normalized["heading"] = expected_heading
    normalized["level"] = int(normalized.get("level") or infer_section_level(expected_heading))
    expected_level = infer_section_level(expected_heading)
    if normalized["level"] != expected_level:
        raise ValueError(f"{expected_heading} 的 level 应为 {expected_level}，实际为 {normalized['level']}。")
    paragraphs = normalized.get("paragraphs")
    if isinstance(paragraphs, str):
        paragraphs = [paragraphs]
    if not isinstance(paragraphs, list):
        raise ValueError(f"{expected_heading} 的 paragraphs 必须是数组。")
    normalized["paragraphs"] = [str(item).strip() for item in paragraphs if str(item).strip()]
    tables = normalized.get("tables", [])
    if tables is None:
        tables = []
    if not isinstance(tables, list):
        raise ValueError(f"{expected_heading} 的 tables 必须是数组。")
    if expected_heading.startswith("附表") and not tables:
        raise ValueError(f"{expected_heading} 必须包含 tables。")
    for table in tables:
        validate_table_payload(table, expected_heading)
    normalized["tables"] = tables
    return normalized


def allows_empty_kimi_parent_paragraphs(heading: str) -> bool:
    try:
        index = REQUIRED_SECTION_HEADINGS.index(heading)
    except ValueError:
        return False
    if index >= len(REQUIRED_SECTION_HEADINGS) - 1:
        return False
    current_level = infer_section_level(heading)
    next_level = infer_section_level(REQUIRED_SECTION_HEADINGS[index + 1])
    return next_level > current_level


def validate_kimi_sections_fragment(data: dict, headings_batch: list[str]) -> list[dict]:
    sections = data.get("sections")
    if not isinstance(sections, list):
        raise ValueError("section 片段缺少 sections 数组。")
    if len(sections) != len(headings_batch):
        returned_headings = [
            str(section.get("heading", "")).strip() if isinstance(section, dict) else "<非对象>"
            for section in sections
        ]
        missing = [heading for heading in headings_batch if heading not in returned_headings]
        raise ValueError(
            f"section 片段章节数量不匹配：期望 {len(headings_batch)} 个，实际 {len(sections)} 个。"
            + (" 缺失：" + "、".join(missing) if missing else "")
        )
    normalized_sections: list[dict] = []
    for section, heading in zip(sections, headings_batch):
        normalized = normalize_kimi_section(section, heading)
        if not normalized["paragraphs"] and not allows_empty_kimi_parent_paragraphs(heading):
            raise ValueError(f"{heading} 缺少有效 paragraphs。")
        normalized_sections.append(normalized)
    return normalized_sections


def call_kimi_json_fragment(
    base_url: str,
    api_key: str,
    model: str,
    prompt: str,
    stage_name: str,
    validator,
):
    active_prompt = prompt
    last_error: Exception | None = None
    for attempt in range(KIMI_SEGMENT_RETRY_LIMIT + 1):
        model_output = call_kimi_fragment_model(base_url, api_key, model, active_prompt)
        try:
            data = extract_json_object(model_output)
            return validator(data)
        except (json.JSONDecodeError, ValueError) as exc:
            last_error = exc
            if attempt >= KIMI_SEGMENT_RETRY_LIMIT:
                excerpt = model_output_excerpt(model_output)
                raise ValueError(
                    f"Kimi 分段生成失败：{stage_name}。原因：{exc}\n"
                    f"模型原始片段摘录（最多前后各 1000 字）：\n{excerpt}"
                ) from exc
            active_prompt = build_kimi_segment_retry_prompt(prompt, stage_name, str(exc))
    if last_error:
        raise last_error
    raise ValueError(f"Kimi 分段生成失败：{stage_name}。")


def validate_kimi_assembled_protocol(data: dict) -> None:
    validate_protocol_json(data)
    headings = [
        str(section.get("heading", "")).strip()
        for section in data.get("sections", [])
        if isinstance(section, dict)
    ]
    if headings != REQUIRED_SECTION_HEADINGS:
        missing = [heading for heading in REQUIRED_SECTION_HEADINGS if heading not in headings]
        extras = [heading for heading in headings if heading not in REQUIRED_SECTION_HEADINGS]
        detail = []
        if missing:
            detail.append("缺失：" + "、".join(missing))
        if extras:
            detail.append("多余：" + "、".join(extras))
        raise ValueError("Kimi 分段组装后的章节顺序或标题不完整。" + "；".join(detail))


def generate_protocol_json_with_kimi_batches(
    base_url: str,
    api_key: str,
    model: str,
    summary_text: str,
) -> dict:
    cover_prompt = build_kimi_cover_summary_prompt(summary_text)
    cover, summary_rows = call_kimi_json_fragment(
        base_url,
        api_key,
        model,
        cover_prompt,
        "cover_summary",
        validate_kimi_cover_summary_fragment,
    )
    sections: list[dict] = []
    generated_headings: list[str] = []
    for index, headings_batch in enumerate(KIMI_SECTION_BATCHES, start=1):
        first_heading = headings_batch[0]
        last_heading = headings_batch[-1]
        stage_name = f"sections_batch_{index}（{first_heading} 至 {last_heading}）"
        section_prompt = build_kimi_sections_prompt(
            summary_text,
            headings_batch,
            cover,
            summary_rows,
            generated_headings,
        )
        batch_sections = call_kimi_json_fragment(
            base_url,
            api_key,
            model,
            section_prompt,
            stage_name,
            lambda data, batch=headings_batch: validate_kimi_sections_fragment(data, batch),
        )
        sections.extend(batch_sections)
        generated_headings.extend(headings_batch)
    protocol_data = {"cover": cover, "summary": summary_rows, "sections": sections}
    validate_kimi_assembled_protocol(protocol_data)
    return protocol_data


def generate_protocol_json_with_validation_retry(prompt: str, caller, max_attempts: int = 1) -> dict:
    active_prompt = prompt
    last_error: Exception | None = None
    for attempt in range(max_attempts):
        model_output = caller(active_prompt)
        try:
            protocol_data = extract_json_object(model_output)
            validate_protocol_json(protocol_data)
            return protocol_data
        except (json.JSONDecodeError, ValueError) as exc:
            last_error = exc
            if attempt >= max_attempts - 1:
                raise
            active_prompt = build_validation_retry_prompt(prompt, str(exc))
    if last_error:
        raise last_error
    raise ValueError("模型输出未能形成完整方案 JSON。")


def generate_protocol_json_deepseek(base_url: str, api_key: str, model: str, prompt: str) -> dict:
    return generate_protocol_json_with_validation_retry(
        prompt,
        lambda active_prompt: call_deepseek_model(base_url, api_key, model, active_prompt),
    )


def generate_protocol_json_kimi(
    base_url: str,
    api_key: str,
    model: str,
    summary_text: str | None,
) -> dict:
    if not summary_text:
        raise ValueError("Kimi 分段生成需要有效的方案摘要文本。")
    return generate_protocol_json_with_kimi_batches(base_url, api_key, model, summary_text)


def generate_protocol_json_glm(base_url: str, api_key: str, model: str, prompt: str) -> dict:
    return generate_protocol_json_with_validation_retry(
        prompt,
        lambda active_prompt: call_glm_model(base_url, api_key, model, active_prompt),
    )


def generate_protocol_json_mimo(base_url: str, api_key: str, model: str, prompt: str) -> dict:
    return generate_protocol_json_with_validation_retry(
        prompt,
        lambda active_prompt: call_mimo_model(base_url, api_key, model, active_prompt),
        max_attempts=MIMO_VALIDATION_RETRY_LIMIT + 1,
    )


def generate_protocol_json_custom(base_url: str, api_key: str, model: str, prompt: str) -> dict:
    return generate_protocol_json_with_validation_retry(
        prompt,
        lambda active_prompt: call_custom_model(base_url, api_key, model, active_prompt),
    )


def generate_protocol_json(
    base_url: str,
    api_key: str,
    model: str,
    prompt: str,
    provider: str,
    summary_text: str | None = None,
) -> dict:
    if not api_key.strip():
        raise ValueError("请填写 api_key。")
    if not model.strip():
        raise ValueError("请填写模型名称。")
    if provider == "deepseek":
        return generate_protocol_json_deepseek(base_url, api_key, model, prompt)
    if provider == "kimi":
        return generate_protocol_json_kimi(base_url, api_key, model, summary_text)
    if provider == "GLM":
        return generate_protocol_json_glm(base_url, api_key, model, prompt)
    if provider == "mimo":
        return generate_protocol_json_mimo(base_url, api_key, model, prompt)
    return generate_protocol_json_custom(base_url, api_key, model, prompt)


def ensure_docx_name(name: str) -> str:
    cleaned = name.strip() or "cohortprotocol.docx"
    cleaned = re.sub(r'[\\/:*?"<>|]', "_", cleaned)
    if not cleaned.lower().endswith(".docx"):
        cleaned += ".docx"
    stem = Path(cleaned).stem.rstrip(" .") or "cohortprotocol"
    reserved_names = {
        "CON", "PRN", "AUX", "NUL",
        "COM1", "COM2", "COM3", "COM4", "COM5", "COM6", "COM7", "COM8", "COM9",
        "LPT1", "LPT2", "LPT3", "LPT4", "LPT5", "LPT6", "LPT7", "LPT8", "LPT9",
    }
    if stem.upper() in reserved_names:
        stem = f"{stem}_"
    return f"{stem}.docx"


def write_brief(
    brief_path: Path,
    input_path: Path,
    output_path: Path,
    provider: str,
    model: str,
    elapsed: float,
    protocol_data: dict,
) -> None:
    cover = protocol_data.get("cover", {})
    sections = protocol_data.get("sections", [])
    tables = sum(len(section.get("tables", [])) for section in sections if isinstance(section, dict))
    lines = [
        f"{APP_NAME} 运行完成简报",
        f"软件制作方：{PRODUCER}",
        f"作者联系方式：Email: {AUTHOR_EMAIL}",
        f"GitHub：{AUTHOR_GITHUB}",
        f"版本号：{APP_VERSION}",
        f"完成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        f"输入摘要：{input_path}",
        f"输出方案：{output_path}",
        f"功能模块：{ACTIVE_MODULE}",
        f"模型类型：{provider}",
        f"模型名称：{model}",
        THINKING_MODE_NOTICE,
        f"运行耗时：{elapsed:.1f} 秒",
        f"研究题目：{cover.get('研究题目') or cover.get('题目') or '未识别'}",
        f"方案章节数：{len(sections)}",
        f"生成表格数：{tables}",
        "",
        "质量检查摘要：",
        "- 已按 cohort-protocol 技能 schema 生成 cover、summary、sections。",
        "- 已调用内置 DOCX 渲染器套用 corhortCRU-small.docx 模板版式。",
        "- 已执行模板痕迹与占位符校验。",
        "- Word 目录字段在打开文档时可由 Word 自动更新。",
        "",
        "风险提示：",
        RISK_NOTICE,
    ]
    brief_path.write_text("\n".join(lines), encoding="utf-8")


class ProtocolDesignApp:
    def __init__(self, root: Tk) -> None:
        self.root = root
        self._layout_mode: str | None = None
        self._layout_after_id: str | None = None
        self._hint_labels: list[ttk.Label] = []
        self._notice_labels: list[ttk.Label] = []
        self.root.title(f"{APP_NAME} {APP_VERSION}")
        self._configure_window()
        self.root.protocol("WM_DELETE_WINDOW", self.on_close_request)

        self.input_path = StringVar()
        self.module = StringVar(value=ACTIVE_MODULE)
        self.provider = StringVar(value="deepseek")
        self.base_url = StringVar(value=PROVIDER_DEFAULTS["deepseek"]["base_url"])
        self.api_key = StringVar(value=saved_api_key("deepseek"))
        self.model = StringVar(value=PROVIDER_DEFAULTS["deepseek"]["model"])
        self.output_dir = StringVar(value=str(app_root()))
        self.output_name = StringVar(value="cohortprotocol.docx")
        self.show_key = BooleanVar(value=False)
        self.remember_key = BooleanVar(value=has_saved_api_key("deepseek"))
        self.running = False

        self.build_ui()

    def _configure_window(self) -> None:
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        width = min(max(1040, int(screen_width * 0.80)), 1480)
        height = min(max(720, int(screen_height * 0.84)), 960)
        pos_x = max(0, (screen_width - width) // 2)
        pos_y = max(0, (screen_height - height) // 2)
        self.root.geometry(f"{width}x{height}+{pos_x}+{pos_y}")
        self.root.minsize(900, 660)

    def build_ui(self) -> None:
        self._init_styles()
        outer = ttk.Frame(self.root, padding=16, style="App.TFrame")
        outer.pack(fill=BOTH, expand=True)
        self._build_shell(outer)
        self.root.bind("<Configure>", self._on_root_configure, add="+")
        self.root.after_idle(self._refresh_layout)
        self.write_log("请依次选择摘要文档、填写模型接口信息和输出位置，然后点击“确定生成方案”。")

    def _init_styles(self) -> None:
        style = ttk.Style()
        style.theme_use("clam")
        style.configure("App.TFrame", background="#F7FAFC")
        style.configure("Surface.TFrame", background="#FFFFFF")
        style.configure("Footer.TFrame", background="#F7FAFC")
        style.configure(
            "Card.TLabelframe",
            background="#FFFFFF",
            borderwidth=1,
            relief="solid",
            bordercolor="#D6E2EE",
            lightcolor="#D6E2EE",
            darkcolor="#D6E2EE",
        )
        style.configure(
            "Card.TLabelframe.Label",
            background="#FFFFFF",
            foreground="#1F4E79",
            font=("Microsoft YaHei UI", 11, "bold"),
        )
        style.configure(
            "Risk.TLabelframe",
            background="#FFF7ED",
            borderwidth=1,
            relief="solid",
            bordercolor="#F3D6B3",
            lightcolor="#F3D6B3",
            darkcolor="#F3D6B3",
        )
        style.configure(
            "Risk.TLabelframe.Label",
            background="#FFF7ED",
            foreground="#9A3412",
            font=("Microsoft YaHei UI", 11, "bold"),
        )
        style.configure(
            "Info.TLabelframe",
            background="#EAF3FA",
            borderwidth=1,
            relief="solid",
            bordercolor="#C8DBEA",
            lightcolor="#C8DBEA",
            darkcolor="#C8DBEA",
        )
        style.configure(
            "Info.TLabelframe.Label",
            background="#EAF3FA",
            foreground="#1F4E79",
            font=("Microsoft YaHei UI", 11, "bold"),
        )
        style.configure("Title.TLabel", background="#FFFFFF", font=("Microsoft YaHei UI", 18, "bold"), foreground="#1F4E79")
        style.configure("TitleCn.TLabel", background="#FFFFFF", font=("Microsoft YaHei UI", 11), foreground="#52606D")
        style.configure("Meta.TLabel", background="#FFFFFF", font=("Microsoft YaHei UI", 9), foreground="#52606D")
        style.configure("FieldLabel.TLabel", background="#FFFFFF", font=("Microsoft YaHei UI", 10, "bold"), foreground="#102A43")
        style.configure("Body.TLabel", background="#FFFFFF", font=("Microsoft YaHei UI", 10), foreground="#102A43")
        style.configure("Hint.TLabel", background="#FFFFFF", font=("Microsoft YaHei UI", 9), foreground="#52606D")
        style.configure("RiskBody.TLabel", background="#FFF7ED", font=("Microsoft YaHei UI", 10), foreground="#9A3412")
        style.configure("InfoBody.TLabel", background="#EAF3FA", font=("Microsoft YaHei UI", 10), foreground="#1F4E79")
        style.configure("InlineInfo.TFrame", background="#EAF3FA")
        style.configure("InlineInfo.TLabel", background="#EAF3FA", font=("Microsoft YaHei UI", 9), foreground="#1F4E79")
        style.configure("Surface.TCheckbutton", background="#FFFFFF", font=("Microsoft YaHei UI", 9), foreground="#102A43")
        style.map("Surface.TCheckbutton", background=[("active", "#FFFFFF"), ("!disabled", "#FFFFFF")])
        style.configure(
            "Primary.TButton",
            font=("Microsoft YaHei UI", 11, "bold"),
            padding=(18, 10),
            foreground="#FFFFFF",
            background="#1F4E79",
            borderwidth=0,
        )
        style.map(
            "Primary.TButton",
            foreground=[("disabled", "#EEF4F8"), ("!disabled", "#FFFFFF")],
            background=[("disabled", "#9FB4C8"), ("pressed", "#163B5C"), ("active", "#17456C"), ("!disabled", "#1F4E79")],
        )
        style.configure(
            "Secondary.TButton",
            font=("Microsoft YaHei UI", 10),
            padding=(14, 10),
            foreground="#1F4E79",
            background="#FFFFFF",
            borderwidth=1,
        )
        style.map(
            "Secondary.TButton",
            foreground=[("disabled", "#94A3B8"), ("!disabled", "#1F4E79")],
            background=[("pressed", "#E2ECF4"), ("active", "#EDF4F8"), ("!disabled", "#FFFFFF")],
        )

    def _build_shell(self, outer: ttk.Frame) -> None:
        scroll_container = ttk.Frame(outer, style="App.TFrame")
        scroll_container.pack(fill=BOTH, expand=True)

        self._scroll_canvas = Canvas(scroll_container, highlightthickness=0, borderwidth=0, background="#F7FAFC")
        v_scroll = ttk.Scrollbar(scroll_container, orient="vertical", command=self._scroll_canvas.yview)
        self._scroll_canvas.configure(yscrollcommand=v_scroll.set)

        v_scroll.pack(side=RIGHT, fill=Y)
        self._scroll_canvas.pack(side=LEFT, fill=BOTH, expand=True)

        self._scroll_inner = ttk.Frame(self._scroll_canvas, style="App.TFrame", padding=(0, 0, 0, 8))
        self._scroll_window_id = self._scroll_canvas.create_window((0, 0), window=self._scroll_inner, anchor="nw")

        self._scroll_inner.bind("<Configure>", self._on_scroll_inner_configure)
        self._scroll_canvas.bind("<Configure>", self._on_scroll_canvas_configure)
        self._scroll_canvas.bind("<Enter>", self._bind_mousewheel)
        self._scroll_canvas.bind("<Leave>", self._unbind_mousewheel)

        self._build_header_card(self._scroll_inner)

        self.workspace = ttk.Frame(self._scroll_inner, style="App.TFrame")
        self.workspace.pack(fill=BOTH, expand=True, pady=(12, 0))
        self.workspace.columnconfigure(0, weight=1)

        self.left_column = ttk.Frame(self.workspace, style="App.TFrame")
        self.right_column = ttk.Frame(self.workspace, style="App.TFrame")
        self._build_form_cards(self.left_column)
        self._build_notice_cards(self.right_column)

        self.bottom = ttk.Frame(outer, style="Footer.TFrame", padding=(0, 12, 0, 0))
        self.bottom.pack(fill=X)
        self._build_action_bar(self.bottom)
        self._build_status_panel(self.bottom)

    def _build_header_card(self, parent: ttk.Frame) -> None:
        self.header_card = ttk.Frame(parent, style="Surface.TFrame", padding=(18, 16))
        self.header_card.pack(fill=X)
        self.header_card.columnconfigure(0, weight=1)

        self.title_group = ttk.Frame(self.header_card, style="Surface.TFrame")
        self.meta_group = ttk.Frame(self.header_card, style="Surface.TFrame")

        ttk.Label(self.title_group, text=APP_NAME, style="Title.TLabel").pack(anchor="w")
        ttk.Label(self.title_group, text=APP_NAME_EN, style="TitleCn.TLabel").pack(anchor="w", pady=(4, 0))

        ttk.Label(self.meta_group, text=f"{PRODUCER}  |  {APP_VERSION}", style="Meta.TLabel").pack(anchor="e")
        ttk.Label(self.meta_group, text=f"作者联系方式：Email: {AUTHOR_EMAIL}", style="Meta.TLabel").pack(anchor="e", pady=(3, 0))
        ttk.Label(self.meta_group, text=f"GitHub: {AUTHOR_GITHUB}", style="Meta.TLabel").pack(anchor="e", pady=(3, 0))

        ttk.Separator(self.header_card, orient="horizontal").grid(row=2, column=0, columnspan=2, sticky="ew", pady=(14, 0))

    def _build_notice_cards(self, parent: ttk.Frame) -> None:
        self._create_notice_card(parent, "风险提示", RISK_NOTICE, "Risk.TLabelframe", "RiskBody.TLabel")
        self._create_notice_card(parent, "适用范围", SCOPE_NOTICE, "Info.TLabelframe", "InfoBody.TLabel")

    def _create_notice_card(
        self,
        parent: ttk.Frame,
        title: str,
        text: str,
        frame_style: str,
        label_style: str,
    ) -> ttk.Label:
        card = ttk.LabelFrame(parent, text=title, style=frame_style, padding=(14, 12))
        card.pack(fill=X, pady=(0, 12))
        body = ttk.Label(card, text=text, style=label_style, justify=LEFT, wraplength=320)
        body.pack(fill=X)
        self._notice_labels.append(body)
        return body

    def _build_form_cards(self, parent: ttk.Frame) -> None:
        self._build_input_card(parent)
        self._build_model_card(parent)
        self._build_output_card(parent)

    def _build_input_card(self, parent: ttk.Frame) -> None:
        body = self._create_card(parent, "步骤 1｜输入摘要")

        module_row, _, _ = self._add_field_block(
            body,
            0,
            "功能模块",
            "当前仅“回顾性队列研究方案”可用，其余模块预留升级",
        )
        self.module_box = ttk.Combobox(module_row, textvariable=self.module, values=MODULE_OPTIONS, state="readonly")
        self.module_box.grid(row=0, column=0, sticky="ew")
        self.module_box.bind("<<ComboboxSelected>>", self.on_module_change)

        path_row, _, _ = self._add_field_block(body, 1, "方案摘要文档", "支持 .docx / .md / .txt")
        self.input_entry = ttk.Entry(path_row, textvariable=self.input_path)
        self.input_entry.grid(row=0, column=0, sticky="ew")
        self.input_browse_button = ttk.Button(path_row, text="浏览...", style="Secondary.TButton", command=self.pick_input)
        self.input_browse_button.grid(
            row=0,
            column=1,
            padx=(10, 0),
        )

    def _build_model_card(self, parent: ttk.Frame) -> None:
        body = self._create_card(parent, "步骤 2｜模型配置")

        provider_row, _, _ = self._add_field_block(
            body,
            0,
            "大模型类型",
            "选择类型后仍可手动修改 base_url 与模型名",
        )
        self.provider_box = ttk.Combobox(
            provider_row,
            textvariable=self.provider,
            values=list(PROVIDER_DEFAULTS.keys()),
            state="readonly",
        )
        self.provider_box.grid(row=0, column=0, sticky="ew")
        self.provider_box.bind("<<ComboboxSelected>>", self.on_provider_change)

        self._add_inline_note(body, 1, THINKING_MODE_NOTICE)

        base_url_row, _, _ = self._add_field_block(body, 2, "base_url")
        ttk.Entry(base_url_row, textvariable=self.base_url).grid(row=0, column=0, sticky="ew")

        api_key_row, api_key_block, _ = self._add_field_block(body, 3, "api_key")
        self.key_entry = ttk.Entry(api_key_row, textvariable=self.api_key, show="*")
        self.key_entry.grid(row=0, column=0, sticky="ew")
        key_options = ttk.Frame(api_key_block, style="Surface.TFrame")
        key_options.grid(row=2, column=0, sticky="w", pady=(6, 0))
        ttk.Checkbutton(
            key_options,
            text="显示 api_key",
            style="Surface.TCheckbutton",
            variable=self.show_key,
            command=self.toggle_key,
        ).pack(side=LEFT)
        ttk.Checkbutton(
            key_options,
            text="本地保存 api_key",
            style="Surface.TCheckbutton",
            variable=self.remember_key,
            command=self.on_remember_key_change,
        ).pack(side=LEFT, padx=(12, 0))

        model_row, _, _ = self._add_field_block(
            body,
            4,
            "模型名称",
            "可下拉选择，也可手动输入平台实际模型名",
        )
        self.model_box = ttk.Combobox(
            model_row,
            textvariable=self.model,
            values=PROVIDER_MODELS.get(self.provider.get(), []),
            state="normal",
        )
        self.model_box.grid(row=0, column=0, sticky="ew")

    def _build_output_card(self, parent: ttk.Frame) -> None:
        body = self._create_card(parent, "步骤 3｜输出设置")

        output_dir_row, _, _ = self._add_field_block(body, 0, "输出目录")
        self.output_dir_entry = ttk.Entry(output_dir_row, textvariable=self.output_dir)
        self.output_dir_entry.grid(row=0, column=0, sticky="ew")
        self.output_path_button = ttk.Button(output_dir_row, text="浏览...", style="Secondary.TButton", command=self.pick_output_dir)
        self.output_path_button.grid(
            row=0,
            column=1,
            padx=(10, 0),
        )

        output_name_row, _, _ = self._add_field_block(body, 1, "输出文件名")
        self.output_name_entry = ttk.Entry(output_name_row, textvariable=self.output_name)
        self.output_name_entry.grid(row=0, column=0, sticky="ew")

    def _create_card(self, parent: ttk.Frame, title: str) -> ttk.Frame:
        card = ttk.LabelFrame(parent, text=title, style="Card.TLabelframe", padding=(16, 14))
        card.pack(fill=X, pady=(0, 12))
        body = ttk.Frame(card, style="Surface.TFrame")
        body.pack(fill=X, expand=True)
        body.columnconfigure(0, weight=1)
        return body

    def _add_field_block(
        self,
        parent: ttk.Frame,
        row: int,
        label: str,
        hint: str = "",
    ) -> tuple[ttk.Frame, ttk.Frame, ttk.Label | None]:
        block = ttk.Frame(parent, style="Surface.TFrame")
        block.grid(row=row, column=0, sticky="ew", pady=(0, 12))
        block.columnconfigure(0, weight=1)

        ttk.Label(block, text=label, style="FieldLabel.TLabel").grid(row=0, column=0, sticky="w")

        input_row = ttk.Frame(block, style="Surface.TFrame")
        input_row.grid(row=1, column=0, sticky="ew", pady=(6, 0))
        input_row.columnconfigure(0, weight=1)

        hint_label = None
        if hint:
            hint_label = ttk.Label(block, text=hint, style="Hint.TLabel", justify=LEFT, wraplength=560)
            hint_label.grid(row=2, column=0, sticky="w", pady=(6, 0))
            self._hint_labels.append(hint_label)

        return input_row, block, hint_label

    def _add_inline_note(self, parent: ttk.Frame, row: int, text: str) -> None:
        note = ttk.Frame(parent, style="InlineInfo.TFrame", padding=(12, 10))
        note.grid(row=row, column=0, sticky="ew", pady=(0, 12))
        note.columnconfigure(0, weight=1)
        label = ttk.Label(note, text=text, style="InlineInfo.TLabel", justify=LEFT, wraplength=640)
        label.grid(row=0, column=0, sticky="w")
        self._hint_labels.append(label)

    def _build_action_bar(self, parent: ttk.Frame) -> None:
        self.action_bar = ttk.Frame(parent, style="Footer.TFrame")
        self.action_bar.pack(fill=X)

        self.run_button = ttk.Button(self.action_bar, text="确定生成方案", style="Primary.TButton", command=self.start_generation)
        self.help_button = ttk.Button(self.action_bar, text="打开帮助文档", style="Secondary.TButton", command=self.open_help)
        self.open_output_button = ttk.Button(
            self.action_bar,
            text="打开输出目录",
            style="Secondary.TButton",
            command=self.open_output_dir,
        )
        self.exit_button = ttk.Button(self.action_bar, text="退出", style="Secondary.TButton", command=self.on_close_request)

    def _build_status_panel(self, parent: ttk.Frame) -> None:
        self.progress = ttk.Progressbar(parent, mode="indeterminate")
        self.progress.pack(fill=X, pady=(10, 8))

        log_frame = ttk.LabelFrame(parent, text="运行状态与完成简报", style="Card.TLabelframe", padding=(12, 10))
        log_frame.pack(fill=X)
        log_content = ttk.Frame(log_frame, style="Surface.TFrame")
        log_content.pack(fill=BOTH, expand=True)
        log_content.columnconfigure(0, weight=1)
        log_scrollbar = ttk.Scrollbar(log_content, orient="vertical")
        self.log = Text(
            log_content,
            height=6,
            wrap="word",
            font=("Microsoft YaHei UI", 9),
            background="#FFFFFF",
            foreground="#102A43",
            insertbackground="#102A43",
            highlightthickness=1,
            highlightbackground="#D6E2EE",
            relief="flat",
            padx=8,
            pady=8,
            yscrollcommand=log_scrollbar.set,
        )
        log_scrollbar.configure(command=self.log.yview)
        self.log.grid(row=0, column=0, sticky="nsew")
        log_scrollbar.grid(row=0, column=1, sticky="ns")

    def toggle_key(self) -> None:
        self.key_entry.configure(show="" if self.show_key.get() else "*")

    def _on_scroll_inner_configure(self, _event=None) -> None:
        self._scroll_canvas.configure(scrollregion=self._scroll_canvas.bbox("all"))

    def _on_scroll_canvas_configure(self, event=None) -> None:
        if event is None:
            return
        self._scroll_canvas.itemconfigure(self._scroll_window_id, width=event.width)
        self._schedule_layout_refresh()

    def _on_root_configure(self, _event=None) -> None:
        self._schedule_layout_refresh()

    def _schedule_layout_refresh(self) -> None:
        if self._layout_after_id is not None:
            self.root.after_cancel(self._layout_after_id)
        self._layout_after_id = self.root.after(80, self._refresh_layout)

    def _refresh_layout(self) -> None:
        self._layout_after_id = None
        window_width = self.root.winfo_width()
        if window_width <= 1:
            return

        if window_width < 1100:
            mode = "narrow"
        elif window_width < 1400:
            mode = "medium"
        else:
            mode = "wide"

        if mode != self._layout_mode:
            self._apply_layout_mode(mode)
            self._layout_mode = mode

        self._update_wraplengths()
        self._update_bottom_height()

    def _apply_layout_mode(self, mode: str) -> None:
        self.title_group.grid_forget()
        self.meta_group.grid_forget()
        if mode == "narrow":
            self.title_group.grid(row=0, column=0, sticky="ew")
            self.meta_group.grid(row=1, column=0, sticky="ew", pady=(12, 0))
            for label in self.meta_group.winfo_children():
                label.pack_configure(anchor="w")
        else:
            self.title_group.grid(row=0, column=0, sticky="nw")
            self.meta_group.grid(row=0, column=1, sticky="ne", padx=(24, 0))
            for label in self.meta_group.winfo_children():
                label.pack_configure(anchor="e")

        self.left_column.grid_forget()
        self.right_column.grid_forget()
        self.workspace.columnconfigure(0, weight=1)
        self.workspace.columnconfigure(1, weight=0, minsize=0)

        if mode == "wide":
            self.workspace.columnconfigure(1, minsize=360)
            self.right_column.configure(width=360)
            self.right_column.grid_propagate(False)
            self.left_column.grid(row=0, column=0, sticky="nsew", padx=(0, 14))
            self.right_column.grid(row=0, column=1, sticky="new")
        elif mode == "medium":
            self.workspace.columnconfigure(1, minsize=320)
            self.right_column.configure(width=320)
            self.right_column.grid_propagate(False)
            self.left_column.grid(row=0, column=0, sticky="nsew", padx=(0, 12))
            self.right_column.grid(row=0, column=1, sticky="new")
        else:
            self.right_column.grid_propagate(True)
            self.right_column.configure(width=0)
            self.left_column.grid(row=0, column=0, sticky="ew")
            self.right_column.grid(row=1, column=0, sticky="ew", pady=(12, 0))

        self._arrange_action_buttons(mode)

    def _arrange_action_buttons(self, mode: str) -> None:
        for button in (self.run_button, self.help_button, self.open_output_button, self.exit_button):
            button.grid_forget()

        for column in range(5):
            self.action_bar.columnconfigure(column, weight=0, minsize=0)

        if mode == "narrow":
            self.action_bar.columnconfigure(0, weight=1)
            self.action_bar.columnconfigure(1, weight=1)
            self.run_button.grid(row=0, column=0, sticky="ew", padx=(0, 8), pady=(0, 8))
            self.help_button.grid(row=0, column=1, sticky="ew", pady=(0, 8))
            self.open_output_button.grid(row=1, column=0, sticky="ew", padx=(0, 8))
            self.exit_button.grid(row=1, column=1, sticky="ew")
        else:
            self.action_bar.columnconfigure(3, weight=1)
            self.run_button.grid(row=0, column=0, sticky="w")
            self.help_button.grid(row=0, column=1, sticky="w", padx=(10, 0))
            self.open_output_button.grid(row=0, column=2, sticky="w", padx=(10, 0))
            self.exit_button.grid(row=0, column=4, sticky="e")

    def _update_wraplengths(self) -> None:
        for label in self._notice_labels:
            available_width = label.master.winfo_width() - 20
            if self._layout_mode == "wide":
                label.configure(wraplength=320)
            elif self._layout_mode == "medium":
                label.configure(wraplength=300)
            elif available_width > 120:
                label.configure(wraplength=max(320, available_width))

        for label in self._hint_labels:
            available_width = label.master.winfo_width() - 12
            if available_width > 80:
                label.configure(wraplength=max(220, min(680, available_width)))

    def _update_bottom_height(self) -> None:
        if not hasattr(self, "log") or not self.root.winfo_exists():
            return
        window_height = self.root.winfo_height()
        if window_height <= 1:
            return
        if window_height < 780:
            lines = 6
        elif window_height < 920:
            lines = 7
        else:
            lines = 8
        if int(self.log.cget("height")) != lines:
            self.log.configure(height=lines)

    def _bind_mousewheel(self, _event=None) -> None:
        self._scroll_canvas.bind_all("<MouseWheel>", self._on_mousewheel)
        self._scroll_canvas.bind_all("<Button-4>", self._on_mousewheel_linux)
        self._scroll_canvas.bind_all("<Button-5>", self._on_mousewheel_linux)

    def _unbind_mousewheel(self, _event=None) -> None:
        self._scroll_canvas.unbind_all("<MouseWheel>")
        self._scroll_canvas.unbind_all("<Button-4>")
        self._scroll_canvas.unbind_all("<Button-5>")

    def _on_mousewheel(self, event) -> None:
        self._scroll_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def _on_mousewheel_linux(self, event) -> None:
        if event.num == 4:
            self._scroll_canvas.yview_scroll(-1, "units")
        elif event.num == 5:
            self._scroll_canvas.yview_scroll(1, "units")

    def on_module_change(self, _event=None) -> None:
        if self.module.get() != ACTIVE_MODULE:
            messagebox.showinfo("模块待开发", "该模块为后续升级预留，当前版本仅支持“回顾性队列研究方案”。")

    def on_remember_key_change(self) -> None:
        provider = self.provider.get()
        if self.remember_key.get():
            api_key = self.api_key.get().strip()
            if api_key:
                try:
                    save_api_key(provider, api_key)
                    self.write_log(f"已使用 Windows 当前用户加密方式保存 {provider} 的 api_key。")
                except Exception as exc:
                    self.remember_key.set(False)
                    messagebox.showerror("保存失败", f"api_key 加密保存失败：{exc}")
        else:
            try:
                delete_saved_api_key(provider)
                self.write_log(f"已删除 {provider} 的本地 api_key 保存项。")
            except Exception as exc:
                self.remember_key.set(True)
                messagebox.showerror("删除失败", f"api_key 本地保存项删除失败：{exc}")

    def sync_api_key_storage(self, provider: str, api_key: str, remember_key: bool) -> None:
        if remember_key:
            if api_key:
                save_api_key(provider, api_key)
        else:
            delete_saved_api_key(provider)

    def build_generation_workflow(self) -> GenerationWorkflow:
        return GenerationWorkflow(
            module=self.module.get(),
            provider=self.provider.get(),
            base_url=self.base_url.get().strip(),
            api_key=self.api_key.get().strip(),
            model=self.model.get().strip(),
            input_path=Path(self.input_path.get().strip()),
            output_dir=Path(self.output_dir.get().strip() or app_root()),
            output_name=ensure_docx_name(self.output_name.get()),
            remember_key=self.remember_key.get(),
        )

    def on_provider_change(self, _event=None) -> None:
        provider = self.provider.get()
        defaults = PROVIDER_DEFAULTS.get(provider, {})
        self.base_url.set(defaults.get("base_url", ""))
        self.model.set(defaults.get("model", ""))
        loaded_key = saved_api_key(provider)
        self.api_key.set(loaded_key)
        self.remember_key.set(bool(loaded_key))
        if hasattr(self, "model_box"):
            self.model_box.configure(values=PROVIDER_MODELS.get(provider, []))

    def pick_input(self) -> None:
        path = filedialog.askopenfilename(
            title="选择方案摘要文档",
            filetypes=[
                ("支持的摘要文档", "*.docx *.md *.markdown *.txt"),
                ("Word 文档", "*.docx"),
                ("Markdown 文档", "*.md *.markdown"),
                ("文本文件", "*.txt"),
                ("所有文件", "*.*"),
            ],
        )
        if path:
            self.input_path.set(path)
            if self.output_name.get() == "cohortprotocol.docx":
                self.output_name.set(f"{Path(path).stem}_完整研究方案.docx")

    def pick_output_dir(self) -> None:
        path = filedialog.askdirectory(title="选择输出目录")
        if path:
            self.output_dir.set(path)

    def open_help(self) -> None:
        project_root = Path(__file__).resolve().parents[1]
        candidates = [
            app_root() / "ProtocolDesign_帮助文档.docx",
            app_root() / "ProtocolDesign_帮助文档.pdf",
            project_root / "ProtocolDesign_帮助文档.docx",
            project_root / "ProtocolDesign_帮助文档.pdf",
            project_root / "docs" / "ProtocolDesign_帮助文档.docx",
            project_root / "docs" / "ProtocolDesign_帮助文档.pdf",
            project_root / "docs" / "ProtocolDesign_帮助文档.md",
            app_root() / "ProtocolDesign_帮助文档.md",
        ]
        seen: set[Path] = set()
        last_error: OSError | None = None
        for help_path in candidates:
            if help_path in seen:
                continue
            seen.add(help_path)
            if not help_path.exists():
                continue
            try:
                os.startfile(str(help_path))
                return
            except OSError as exc:
                last_error = exc
        if last_error is not None:
            messagebox.showerror("帮助文档", f"帮助文档存在但打开失败：{last_error}")
            return
        messagebox.showinfo("帮助文档", "帮助文档未找到，请确认程序目录完整。")

    def open_output_dir(self) -> None:
        path = Path(self.output_dir.get().strip() or app_root())
        try:
            path.mkdir(parents=True, exist_ok=True)
            os.startfile(str(path))
        except OSError as exc:
            messagebox.showerror("打开失败", f"无法打开输出目录：{exc}")

    def ui_call(self, callback, *args) -> None:
        try:
            if not self.root.winfo_exists():
                return
            self.root.after(0, callback, *args)
        except TclError:
            return

    def on_close_request(self) -> None:
        if self.running:
            should_close = messagebox.askyesno(
                "确认退出",
                "当前正在生成方案，立即退出会中断本次生成。\n\n确定要退出吗？",
            )
            if not should_close:
                return
        self.root.destroy()

    def write_log(self, text: str) -> None:
        stamp = datetime.now().strftime("%H:%M:%S")
        self.log.insert(END, f"[{stamp}] {text}\n")
        self.log.see(END)
        self.root.update_idletasks()

    def set_running(self, value: bool) -> None:
        self.running = value
        self.run_button.configure(state="disabled" if value else "normal")
        if value:
            self.progress.start(10)
        else:
            self.progress.stop()

    def start_generation(self) -> None:
        if self.running:
            return
        workflow = self.build_generation_workflow()
        try:
            self.sync_api_key_storage(workflow.provider, workflow.api_key, workflow.remember_key)
        except Exception as exc:
            self.write_log(f"提示：本地保存 api_key 失败，本次生成仍将继续。原因：{exc}")
        self.set_running(True)
        worker = threading.Thread(target=self.generate, args=(workflow,), daemon=True)
        worker.start()

    def generate(self, workflow: GenerationWorkflow) -> None:
        started = time.time()
        try:
            if workflow.module != ACTIVE_MODULE:
                raise ValueError("当前所选模块尚未开发。请先选择“回顾性队列研究方案”。")
            if not workflow.input_path.exists():
                raise ValueError("请选择有效的方案摘要文档。")
            output_path = workflow.output_dir / workflow.output_name
            brief_path = output_path.with_name(f"{output_path.stem}_运行简报.txt")

            self.ui_call(self.write_log, "正在读取方案摘要文档...")
            summary_text = read_summary(workflow.input_path)
            if len(summary_text.strip()) < 20:
                raise ValueError("摘要文档内容过少，无法生成完整研究方案。")

            self.ui_call(self.write_log, "正在组装 cohort-protocol 技能提示词与模板写作规则...")
            prompt = build_prompt(summary_text)

            self.ui_call(self.write_log, f"正在调用 {workflow.provider} 独立工作流生成结构化研究方案 JSON；{THINKING_MODE_NOTICE}")
            protocol_data = generate_protocol_json(
                workflow.base_url,
                workflow.api_key,
                workflow.model,
                prompt,
                workflow.provider,
                summary_text,
            )
            self.ui_call(self.write_log, "模型输出已返回，正在解析并校验结构化结果...")
            self.ui_call(self.write_log, "正在套用内置模板渲染 DOCX 文档...")
            protocol_renderer.render_protocol(protocol_data, output_path)

            elapsed = time.time() - started
            write_brief(
                brief_path,
                workflow.input_path,
                output_path,
                workflow.provider,
                workflow.model,
                elapsed,
                protocol_data,
            )
            self.ui_call(self.write_log, f"生成完成：{output_path}")
            self.ui_call(self.write_log, f"运行简报：{brief_path}")
            self.ui_call(messagebox.showinfo, "完成", f"研究方案已生成：\n{output_path}\n\n运行简报已生成：\n{brief_path}")
        except Exception as exc:
            error_path = app_root() / "ProtocolDesign_error.log"
            detail_message = f"详细错误已保存：\n{error_path}"
            try:
                error_path.write_text(traceback.format_exc(), encoding="utf-8")
            except Exception as log_exc:
                detail_message = f"详细错误写入失败：{log_exc}"
            self.ui_call(self.write_log, f"生成失败：{exc}")
            self.ui_call(messagebox.showerror, "生成失败", f"{exc}\n\n{detail_message}")
        finally:
            self.ui_call(self.set_running, False)


def main() -> None:
    root = Tk()
    ProtocolDesignApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
